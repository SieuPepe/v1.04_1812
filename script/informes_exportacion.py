# script/informes_exportacion.py
"""
Módulo de exportación de informes a Excel, Word y PDF
Genera documentos profesionales con agrupaciones, subtotales y formato corporativo
"""

import os
import subprocess
from datetime import datetime, date
from typing import List, Dict, Any, Optional
from decimal import Decimal
import xlsxwriter
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from PIL import Image as PILImage


class InformesExportador:
    """Exportador de informes a múltiples formatos"""

    def __init__(self, schema: str):
        self.schema = schema
        self.logo_redes_path = None
        self.logo_urbide_path = None
        self._buscar_logos()

    @staticmethod
    def _calcular_escala_imagen(ruta_imagen: str, altura_deseada_cm: float = 2.0) -> tuple:
        """
        Calcula la escala necesaria para que una imagen tenga una altura específica en Excel

        Args:
            ruta_imagen: Ruta a la imagen
            altura_deseada_cm: Altura deseada en centímetros (default 2.0cm)

        Returns:
            Tupla (x_scale, y_scale, ancho_escalado_cm, ancho_px, alto_px) para mantener aspect ratio con la altura deseada
        """
        try:
            # Abrir imagen y obtener dimensiones en píxeles
            img = PILImage.open(ruta_imagen)
            ancho_px, alto_px = img.size

            # Obtener DPI de la imagen (por defecto 96 si no está definido)
            dpi = img.info.get('dpi', (96, 96))
            if isinstance(dpi, tuple):
                dpi = dpi[0]  # Usar DPI horizontal

            # Calcular tamaño actual de la imagen en cm
            # px / DPI = pulgadas; pulgadas * 2.54 = cm
            altura_actual_cm = (alto_px / dpi) * 2.54
            ancho_actual_cm = (ancho_px / dpi) * 2.54

            # Calcular escala necesaria
            scale = altura_deseada_cm / altura_actual_cm

            # Calcular ancho resultante en cm
            ancho_escalado_cm = ancho_actual_cm * scale

            print(f"DEBUG Imagen {ruta_imagen.split('/')[-1]}: {ancho_px}x{alto_px}px @ {dpi}DPI")
            print(f"  Tamaño actual: {ancho_actual_cm:.2f}x{altura_actual_cm:.2f}cm")
            print(f"  Scale calculado: {scale:.4f}")
            print(f"  Tamaño final: {ancho_escalado_cm:.2f}x{altura_deseada_cm:.2f}cm")

            return (scale, scale, ancho_escalado_cm, ancho_px, alto_px)
        except Exception as e:
            print(f"Error calculando escala de imagen {ruta_imagen}: {e}")
            import traceback
            traceback.print_exc()
            return (1.0, 1.0, 0, 0, 0)  # Escala por defecto

    def _buscar_logos(self):
        """Busca los logos en la raíz del proyecto y en la carpeta resources/images"""
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

        # Directorios donde buscar (en orden de prioridad)
        directorios_busqueda = [
            os.path.join(base_dir, "resources", "images"),  # Carpeta resources/images (prioridad 1)
            base_dir,  # Raíz del proyecto (prioridad 2)
        ]

        for directorio in directorios_busqueda:
            if not os.path.exists(directorio):
                continue

            archivos = os.listdir(directorio)

            # Buscar logos con prioridad exacta
            for file in archivos:
                file_lower = file.lower()
                if not file_lower.endswith(('.png', '.jpg', '.jpeg')):
                    continue

                # Logo izquierdo (Logo Redes Urbide) - prioridad exacta
                if not self.logo_redes_path:
                    if file_lower in ["logo redes urbide.jpg", "logo redes urbide.png"]:
                        self.logo_redes_path = os.path.join(directorio, file)
                        print(f"✓ Logo izquierdo (Redes Urbide) encontrado: {file}")

                # Logo derecho (Logo Urbide) - prioridad exacta
                if not self.logo_urbide_path:
                    if file_lower in ["logo urbide.jpg", "logo urbide.png"]:
                        self.logo_urbide_path = os.path.join(directorio, file)
                        print(f"✓ Logo derecho (Urbide) encontrado: {file}")

            # Si ya encontramos ambos logos, salir
            if self.logo_redes_path and self.logo_urbide_path:
                break

        # Si aún no se encuentran, buscar alternativas en todos los directorios
        if not self.logo_redes_path or not self.logo_urbide_path:
            for directorio in directorios_busqueda:
                if not os.path.exists(directorio):
                    continue

                archivos = os.listdir(directorio)

                for file in archivos:
                    file_lower = file.lower()
                    if not file_lower.endswith(('.png', '.jpg', '.jpeg')):
                        continue

                    # Intentar con patrones flexibles como fallback
                    if not self.logo_redes_path and ("redes" in file_lower or file_lower == "logo.png"):
                        self.logo_redes_path = os.path.join(directorio, file)
                        print(f"⚠ Usando logo alternativo (izquierdo): {file}")

                    if not self.logo_urbide_path and "urbide" in file_lower and "redes" not in file_lower:
                        self.logo_urbide_path = os.path.join(directorio, file)
                        print(f"⚠ Usando logo alternativo (derecho): {file}")

    def exportar_a_excel(
        self,
        filepath: str,
        informe_nombre: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        proyecto_nombre: str = "",
        proyecto_codigo: str = "",
        fecha_informe: str = ""
    ) -> bool:
        """
        Exporta el informe a Excel con formato profesional

        Args:
            filepath: Ruta del archivo Excel a crear
            informe_nombre: Nombre del informe
            columnas: Lista de nombres de columnas
            datos: Datos del informe
            resultado_agrupacion: Estructura de agrupaciones y totales (opcional)
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto

        Returns:
            True si la exportación fue exitosa
        """
        try:
            workbook = xlsxwriter.Workbook(filepath)
            worksheet = workbook.add_worksheet(informe_nombre[:31])  # Excel limit 31 chars

            # Función helper para detectar y convertir fechas
            def detectar_y_convertir_fecha(valor):
                """Detecta formatos comunes de fecha y los convierte a datetime"""
                import re
                from datetime import datetime as dt, date

                # Si ya es un objeto datetime o date, retornarlo
                if isinstance(valor, (dt, date)):
                    if isinstance(valor, date) and not isinstance(valor, dt):
                        # Convertir date a datetime
                        return dt.combine(valor, dt.min.time())
                    return valor

                # Si es string, detectar formato
                if not isinstance(valor, str):
                    return None

                # Detectar formatos comunes: DD/MM/YYYY, YYYY-MM-DD, DD-MM-YYYY
                patrones_fecha = [
                    (r'^\d{2}/\d{2}/\d{4}$', '%d/%m/%Y'),  # DD/MM/YYYY
                    (r'^\d{4}-\d{2}-\d{2}$', '%Y-%m-%d'),  # YYYY-MM-DD
                    (r'^\d{2}-\d{2}-\d{4}$', '%d-%m-%Y'),  # DD-MM-YYYY
                ]

                for patron, formato in patrones_fecha:
                    if re.match(patron, valor.strip()):
                        try:
                            return dt.strptime(valor.strip(), formato)
                        except ValueError:
                            continue
                return None

            # Definir formatos
            formato_titulo = workbook.add_format({
                'bold': True,
                'font_size': 20,
                'font_name': 'Calibri',
                'align': 'center',  # Centrado horizontal
                'valign': 'vcenter'  # Centrado vertical
            })

            formato_subtitulo = workbook.add_format({
                'font_size': 10,
                'font_name': 'Tahoma',
                'italic': True,
                'font_color': '#7C7C7C',
                'align': 'left'
            })

            formato_header_nivel0 = workbook.add_format({
                'bold': True,
                'font_size': 10,
                'font_name': 'Tahoma',
                'bg_color': '#4A6FA5',
                'font_color': 'white',
                'border': 1,
                'align': 'left',
                'valign': 'vcenter'
            })

            formato_header_nivel1 = workbook.add_format({
                'bold': True,
                'font_size': 9,
                'font_name': 'Tahoma',
                'bg_color': '#6B8FB8',
                'font_color': 'white',
                'border': 1,
                'align': 'left'
            })

            formato_header_nivel2 = workbook.add_format({
                'bold': True,
                'font_size': 9,
                'font_name': 'Tahoma',
                'bg_color': '#8AADC7',
                'font_color': 'white',
                'border': 1,
                'align': 'left'
            })

            formato_header_columnas = workbook.add_format({
                'bold': True,
                'font_size': 10,
                'font_name': 'Tahoma',
                'bg_color': '#D9D9D9',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })

            formato_datos = workbook.add_format({
                'font_size': 8,
                'font_name': 'Tahoma',
                'border': 1
            })

            formato_moneda = workbook.add_format({
                'font_size': 8,
                'font_name': 'Tahoma',
                'num_format': '#,##0.00 €',
                'border': 1,
                'align': 'right',  # Alineación a la derecha
                'valign': 'vcenter'
            })

            formato_decimal = workbook.add_format({
                'font_size': 8,
                'font_name': 'Tahoma',
                'num_format': '#,##0.00',
                'border': 1,
                'align': 'right',  # Alineación a la derecha
                'valign': 'vcenter'
            })

            # Formato específico para coordenadas (latitud/longitud) con 4 decimales
            formato_coordenadas = workbook.add_format({
                'font_size': 8,
                'font_name': 'Tahoma',
                'num_format': '#,##0.0000',
                'border': 1,
                'align': 'right',
                'valign': 'vcenter'
            })

            formato_subtotal = workbook.add_format({
                'bold': True,
                'font_size': 9,
                'font_name': 'Tahoma',
                'bg_color': '#E7E6E6',
                'num_format': '#,##0.00 €',
                'border': 1,
                'align': 'right',  # Alineación a la derecha
                'valign': 'vcenter'
            })

            formato_subtotal_texto = workbook.add_format({
                'bold': True,
                'font_size': 9,
                'font_name': 'Tahoma',
                'bg_color': '#E7E6E6',
                'border': 1,
                'align': 'left',  # Texto alineado a la izquierda
                'valign': 'vcenter'
            })

            formato_total = workbook.add_format({
                'bold': True,
                'font_size': 10,
                'font_name': 'Tahoma',
                'bg_color': '#C5D9F1',
                'num_format': '#,##0.00 €',
                'border': 2,
                'align': 'right',  # Alineación a la derecha
                'valign': 'vcenter'
            })

            formato_total_texto = workbook.add_format({
                'bold': True,
                'font_size': 10,
                'font_name': 'Tahoma',
                'bg_color': '#C5D9F1',
                'border': 2,
                'align': 'left',  # Texto alineado a la izquierda
                'valign': 'vcenter'
            })

            formato_fecha = workbook.add_format({
                'font_size': 10,
                'font_name': 'Calibri',
                'bold': True
            })

            # Formato para celdas de fecha (dd/mm/yyyy)
            formato_fecha_celda = workbook.add_format({
                'font_size': 8,
                'font_name': 'Tahoma',
                'border': 1,
                'num_format': 'dd/mm/yyyy',
                'align': 'left',
                'valign': 'vcenter'
            })

            # Fila actual
            row = 0

            # Configurar altura de la fila de encabezado para acomodar logos y título
            # 2.1 cm ≈ 59.5 puntos
            worksheet.set_row(row, 59.5)

            # Calcular anchos de columnas para los logos dinámicamente
            # Logo Redes Urbide (izquierda)
            ancho_col_izq_chars = 15  # Default
            if self.logo_redes_path and os.path.exists(self.logo_redes_path):
                x_scale_izq, _, ancho_img_cm_izq, ancho_px_izq, _ = self._calcular_escala_imagen(self.logo_redes_path, 2.0)
                # XlsxWriter convierte imágenes a 96 DPI internamente
                ancho_escalado_px_izq = (ancho_img_cm_izq / 2.54) * 96
                # Convertir píxeles a caracteres: chars = (pixels - 5) / 7
                # Añadir margen de 10px (≈1.4 chars) para espacio
                ancho_col_izq_chars = max(15, int((ancho_escalado_px_izq + 10) / 7))

            # Logo Urbide (derecha) - mantener ancho estándar de columna
            ancho_col_der_chars = 15  # Ancho estándar, NO cambiar por la imagen
            ancho_img_px_escalado_derecha = 0  # Para usar más tarde
            if self.logo_urbide_path and os.path.exists(self.logo_urbide_path):
                x_scale_der, _, ancho_img_cm_derecha, ancho_px_der, _ = self._calcular_escala_imagen(self.logo_urbide_path, 2.0)

                # XlsxWriter convierte imágenes a 96 DPI internamente
                # Convertir el ancho en cm a píxeles @ 96 DPI
                ancho_img_px_escalado_derecha = (ancho_img_cm_derecha / 2.54) * 96

                print(f"DEBUG Logo Urbide - Cálculo de escala:")
                print(f"  Ancho imagen en cm: {ancho_img_cm_derecha:.2f}cm")
                print(f"  Ancho en píxeles @ 96 DPI: {ancho_img_px_escalado_derecha:.1f}px")
                print(f"  Ancho columna estándar: {ancho_col_der_chars} chars (sin modificar)")

            # Configurar ancho de primera y última columna
            worksheet.set_column(0, 0, ancho_col_izq_chars)
            # La última columna mantiene el ancho estándar (15 chars)
            worksheet.set_column(len(columnas) - 1, len(columnas) - 1, ancho_col_der_chars)

            # Logo izquierdo (Logo Redes Urbide) - altura 2.0cm, alineado a la izquierda
            if self.logo_redes_path and os.path.exists(self.logo_redes_path):
                x_scale, y_scale, ancho_img_cm, _, _ = self._calcular_escala_imagen(self.logo_redes_path, 2.0)
                worksheet.insert_image(row, 0, self.logo_redes_path, {
                    'x_scale': x_scale,
                    'y_scale': y_scale,
                    'x_offset': 2,
                    'y_offset': 2,
                    'object_position': 1  # Mover con celda y redimensionar
                })

            # Título del informe en el centro (entre los logos)
            # Combinar celdas centrales y escribir el título
            num_cols = len(columnas)
            col_inicio_titulo = 1
            col_fin_titulo = num_cols - 2
            if col_fin_titulo <= col_inicio_titulo:
                col_fin_titulo = col_inicio_titulo

            # Escribir el título combinando celdas
            worksheet.merge_range(row, col_inicio_titulo, row, col_fin_titulo, informe_nombre, formato_titulo)

            # Logo derecho (Logo Urbide) - altura exacta 2cm, alineado al borde derecho de la última columna
            if self.logo_urbide_path and os.path.exists(self.logo_urbide_path) and ancho_img_px_escalado_derecha > 0:
                x_scale, y_scale, _, _, _ = self._calcular_escala_imagen(self.logo_urbide_path, 2.0)

                # Calcular ancho de la última columna en píxeles (ancho estándar 15 chars)
                ancho_ultima_col_px = ancho_col_der_chars * 7 + 5  # 15 * 7 + 5 = 110px

                # Para alinear a la derecha de la última columna:
                # offset = ancho_columna - ancho_imagen_escalada - margen
                # Si la imagen es más ancha que la columna, el offset será negativo (está bien)
                margen_derecho = 3
                x_offset_derecha = ancho_ultima_col_px - ancho_img_px_escalado_derecha - margen_derecho

                print(f"DEBUG Logo Urbide - Posicionamiento con offset (puede ser negativo):")
                print(f"  Ancho última columna: {ancho_col_der_chars} chars → {ancho_ultima_col_px:.1f}px")
                print(f"  Ancho imagen escalada: {ancho_img_px_escalado_derecha:.1f}px")
                print(f"  Margen derecho: {margen_derecho}px")
                print(f"  x_offset = {ancho_ultima_col_px:.1f} - {ancho_img_px_escalado_derecha:.1f} - {margen_derecho} = {x_offset_derecha:.1f}px")

                # Insertar en la ÚLTIMA columna con offset para alinear a la derecha
                # El offset puede ser negativo si la imagen es más ancha que la columna
                worksheet.insert_image(row, len(columnas) - 1, self.logo_urbide_path, {
                    'x_scale': x_scale,
                    'y_scale': y_scale,
                    'x_offset': int(x_offset_derecha),  # Permite offset negativo
                    'y_offset': 2,
                    'object_position': 1  # Mover con celda y redimensionar
                })

            row += 2  # Espacio después del encabezado (1 fila de encabezado + 1 de espacio)

            # Información del proyecto
            if proyecto_nombre:
                worksheet.merge_range(row, 0, row, len(columnas) - 1, proyecto_nombre, formato_subtitulo)
                row += 1

            # Fecha - usar fecha proporcionada por el usuario o generar automáticamente
            fecha_actual = fecha_informe if fecha_informe else datetime.now().strftime("%d/%m/%Y")
            worksheet.write(row, 0, f"FECHA: {fecha_actual}", formato_fecha)
            row += 2

            # Si hay agrupaciones, exportar con estructura jerárquica
            if resultado_agrupacion and resultado_agrupacion.get('grupos'):
                # Si hay columnas_datos, usarlas en lugar de columnas completas
                columnas_para_excel = resultado_agrupacion.get('columnas_datos', columnas)

                row = self._exportar_grupos_excel(
                    worksheet,
                    workbook,
                    resultado_agrupacion['grupos'],
                    columnas_para_excel,  # Usar columnas filtradas
                    row,
                    formato_header_nivel0,
                    formato_header_nivel1,
                    formato_header_nivel2,
                    formato_header_columnas,
                    formato_datos,
                    formato_moneda,
                    formato_decimal,
                    formato_coordenadas,
                    formato_subtotal,
                    formato_subtotal_texto,
                    formato_fecha_celda,
                    resultado_agrupacion.get('modo', 'detalle'),
                    resultado_agrupacion
                )

                # Totales generales - ALINEADOS CON LAS COLUMNAS CORRECTAS
                if resultado_agrupacion.get('totales_generales'):
                    row += 1
                    worksheet.write(row, 0, "═══ TOTAL GENERAL ═══", formato_total_texto)

                    # Obtener mapa de formatos de agregaciones
                    formatos_agregaciones = resultado_agrupacion.get('formatos_agregaciones', {})

                    # Crear formato para total sin moneda (para COUNT, etc.)
                    formato_total_entero = workbook.add_format({
                        'bold': True,
                        'font_size': 10,
                        'font_name': 'Tahoma',
                        'bg_color': '#C5D9F1',
                        'num_format': '#,##0',
                        'border': 2,
                        'align': 'right',
                        'valign': 'vcenter'
                    })

                    totales = resultado_agrupacion['totales_generales']
                    for key, valor in totales.items():
                        # Extraer el nombre del campo del key
                        # Puede ser "SUM(presupuesto)" o directamente "Importe" (nuevo formato)
                        if '(' in key and ')' in key:
                            # Formato antiguo: "FUNCION(campo)"
                            campo_nombre = key.split('(')[1].rstrip(')')
                        else:
                            # Formato nuevo: nombre directo de columna
                            campo_nombre = key

                        # Determinar el formato según el tipo de agregación
                        formato_agg = formatos_agregaciones.get(key, 'ninguno')
                        formato_a_usar = formato_total if formato_agg == 'moneda' else formato_total_entero

                        # Buscar la columna correspondiente en columnas filtradas
                        if campo_nombre in columnas_para_excel:
                            col_idx = columnas_para_excel.index(campo_nombre)
                            worksheet.write(row, col_idx, valor, formato_a_usar)
                        elif campo_nombre == '*':
                            # COUNT(*) se escribe en la segunda columna, siempre como entero
                            worksheet.write(row, 1, valor, formato_total_entero)

            else:
                # Exportar sin agrupaciones (tabla simple)
                # Función helper para detectar y convertir fechas
                def detectar_y_convertir_fecha(valor):
                    """Detecta formatos comunes de fecha y los convierte a datetime"""
                    import re
                    from datetime import datetime as dt, date

                    # Si ya es un objeto datetime o date, retornarlo
                    if isinstance(valor, (dt, date)):
                        if isinstance(valor, date) and not isinstance(valor, dt):
                            # Convertir date a datetime
                            return dt.combine(valor, dt.min.time())
                        return valor

                    # Si es string, detectar formato
                    if not isinstance(valor, str):
                        return None

                    # Detectar formatos comunes: DD/MM/YYYY, YYYY-MM-DD, DD-MM-YYYY
                    patrones_fecha = [
                        (r'^\d{2}/\d{2}/\d{4}$', '%d/%m/%Y'),  # DD/MM/YYYY
                        (r'^\d{4}-\d{2}-\d{2}$', '%Y-%m-%d'),  # YYYY-MM-DD
                        (r'^\d{2}-\d{2}-\d{4}$', '%d-%m-%Y'),  # DD-MM-YYYY
                    ]

                    for patron, formato in patrones_fecha:
                        if re.match(patron, valor.strip()):
                            try:
                                return dt.strptime(valor.strip(), formato)
                            except ValueError:
                                continue
                    return None

                # Encabezados de columnas
                for col_idx, col_name in enumerate(columnas):
                    worksheet.write(row, col_idx, col_name, formato_header_columnas)
                    worksheet.set_column(col_idx, col_idx, 15)  # Ancho por defecto
                row += 1

                # Datos
                formatos_columnas = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}
                for fila_datos in datos:
                    for col_idx, valor in enumerate(fila_datos):
                        # Obtener el formato del campo según su columna
                        col_name = columnas[col_idx] if col_idx < len(columnas) else None
                        formato_campo = formatos_columnas.get(col_name, 'ninguno') if col_name else 'ninguno'

                        # Detectar si es una coordenada (latitud/longitud)
                        es_coordenada = col_name and ('latitud' in col_name.lower() or 'longitud' in col_name.lower())

                        # Detectar y manejar fechas
                        fecha_dt = detectar_y_convertir_fecha(valor)
                        if fecha_dt:
                            worksheet.write_datetime(row, col_idx, fecha_dt, formato_fecha_celda)
                        # Debug: imprimir valores que parecen fechas pero no se detectan
                        elif col_name and 'fecha' in col_name.lower() and valor:
                            print(f"DEBUG Fecha no detectada en columna '{col_name}': {valor} (tipo: {type(valor).__name__})")
                        # Aplicar formato según el tipo de campo
                        elif isinstance(valor, (int, float, Decimal)):
                            # Convertir Decimal a float para poder formatear
                            if isinstance(valor, Decimal):
                                valor = float(valor)

                            if es_coordenada:
                                # Coordenadas: 4 decimales
                                worksheet.write(row, col_idx, valor, formato_coordenadas)
                            elif formato_campo == 'moneda':
                                worksheet.write(row, col_idx, valor, formato_moneda)
                            elif formato_campo == 'decimal':
                                worksheet.write(row, col_idx, valor, formato_decimal)
                            else:
                                # Por defecto, números decimales con 2 decimales
                                worksheet.write(row, col_idx, valor, formato_decimal)
                        else:
                            worksheet.write(row, col_idx, valor, formato_datos)
                    row += 1

            # Pie de página
            row += 2
            worksheet.write(row, 0, f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}", formato_subtitulo)

            workbook.close()
            return True

        except Exception as e:
            print(f"Error al exportar a Excel: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _exportar_grupos_excel(
        self,
        worksheet,
        workbook,
        grupos: List[Dict],
        columnas: List[str],
        start_row: int,
        formato_nivel0,
        formato_nivel1,
        formato_nivel2,
        formato_header_columnas,
        formato_datos,
        formato_moneda,
        formato_decimal,
        formato_coordenadas,
        formato_subtotal,
        formato_subtotal_texto,
        formato_fecha_celda,
        modo: str = 'detalle',
        resultado_agrupacion: Optional[Dict] = None
    ) -> int:
        """Exporta grupos jerárquicos a Excel (recursivo)"""
        # Función helper para detectar y convertir fechas
        def detectar_y_convertir_fecha(valor):
            """Detecta formatos comunes de fecha y los convierte a datetime"""
            import re
            from datetime import datetime as dt, date

            # Si ya es un objeto datetime o date, retornarlo
            if isinstance(valor, (dt, date)):
                if isinstance(valor, date) and not isinstance(valor, dt):
                    # Convertir date a datetime
                    return dt.combine(valor, dt.min.time())
                return valor

            # Si es string, detectar formato
            if not isinstance(valor, str):
                return None

            # Detectar formatos comunes: DD/MM/YYYY, YYYY-MM-DD, DD-MM-YYYY
            patrones_fecha = [
                (r'^\d{2}/\d{2}/\d{4}$', '%d/%m/%Y'),  # DD/MM/YYYY
                (r'^\d{4}-\d{2}-\d{2}$', '%Y-%m-%d'),  # YYYY-MM-DD
                (r'^\d{2}-\d{2}-\d{4}$', '%d-%m-%Y'),  # DD-MM-YYYY
            ]

            for patron, formato in patrones_fecha:
                if re.match(patron, valor.strip()):
                    try:
                        return dt.strptime(valor.strip(), formato)
                    except ValueError:
                        continue
            return None

        row = start_row

        formatos_por_nivel = [formato_nivel0, formato_nivel1, formato_nivel2]

        for grupo in grupos:
            nivel = grupo.get('nivel', 0)
            clave = grupo.get('clave', '')
            campo = grupo.get('campo', '')
            datos = grupo.get('datos', [])
            subtotales = grupo.get('subtotales', {})
            subgrupos = grupo.get('subgrupos')

            formato_grupo = formatos_por_nivel[min(nivel, 2)]

            # Encabezado del grupo
            indent = "    " * nivel
            titulo_grupo = f"{indent}📁 {campo.upper()}: {clave}"
            worksheet.merge_range(row, 0, row, len(columnas) - 1, titulo_grupo, formato_grupo)
            worksheet.set_row(row, 20 + (5 * (2 - nivel)))  # Altura según nivel
            row += 1

            # Si hay subgrupos, procesarlos recursivamente
            if subgrupos:
                row = self._exportar_grupos_excel(
                    worksheet,
                    workbook,
                    subgrupos,
                    columnas,
                    row,
                    formato_nivel0,
                    formato_nivel1,
                    formato_nivel2,
                    formato_header_columnas,
                    formato_datos,
                    formato_moneda,
                    formato_decimal,
                    formato_coordenadas,
                    formato_subtotal,
                    formato_subtotal_texto,
                    formato_fecha_celda,
                    modo,
                    resultado_agrupacion
                )
            elif modo == 'detalle':
                # Encabezados de columnas para este grupo
                for col_idx, col_name in enumerate(columnas):
                    worksheet.write(row, col_idx, col_name, formato_header_columnas)
                row += 1

                # Datos del grupo
                for fila_datos in datos:
                    for col_idx, valor in enumerate(fila_datos):
                        # Obtener el formato del campo según su columna
                        col_name = columnas[col_idx] if col_idx < len(columnas) else None
                        formato_campo = resultado_agrupacion.get('formatos_columnas', {}).get(col_name, 'ninguno') if col_name else 'ninguno'

                        # Detectar si es una coordenada (latitud/longitud)
                        es_coordenada = col_name and ('latitud' in col_name.lower() or 'longitud' in col_name.lower())

                        # Detectar y manejar fechas
                        fecha_dt = detectar_y_convertir_fecha(valor)
                        if fecha_dt:
                            worksheet.write_datetime(row, col_idx, fecha_dt, formato_fecha_celda)
                        # Debug: imprimir valores que parecen fechas pero no se detectan
                        elif col_name and 'fecha' in col_name.lower() and valor:
                            print(f"DEBUG Fecha no detectada en columna '{col_name}': {valor} (tipo: {type(valor).__name__})")
                        # Aplicar formato según el tipo de campo
                        elif isinstance(valor, (int, float, Decimal)):
                            # Convertir Decimal a float para poder formatear
                            if isinstance(valor, Decimal):
                                valor = float(valor)

                            if es_coordenada:
                                # Coordenadas: 4 decimales
                                worksheet.write(row, col_idx, valor, formato_coordenadas)
                            elif formato_campo == 'moneda':
                                worksheet.write(row, col_idx, valor, formato_moneda)
                            elif formato_campo == 'decimal':
                                worksheet.write(row, col_idx, valor, formato_decimal)
                            else:
                                # Por defecto, números decimales con 2 decimales
                                worksheet.write(row, col_idx, valor, formato_decimal)
                        else:
                            worksheet.write(row, col_idx, str(valor) if valor is not None else "", formato_datos)
                    row += 1

            # Subtotales del grupo - ALINEADOS CON LAS COLUMNAS CORRECTAS
            if subtotales:
                indent_subtotal = "    " * (nivel + 1)

                # Escribir "Subtotal" en la primera columna (con formato de texto)
                worksheet.write(row, 0, f"{indent_subtotal}▸ Subtotal", formato_subtotal_texto)

                # Obtener mapa de formatos de agregaciones
                formatos_agregaciones = resultado_agrupacion.get('formatos_agregaciones', {}) if resultado_agrupacion else {}

                # Crear formato para subtotal sin moneda (para COUNT, etc.)
                formato_subtotal_entero = workbook.add_format({
                    'bold': True,
                    'font_size': 9,
                    'font_name': 'Tahoma',
                    'bg_color': '#E7E6E6',
                    'num_format': '#,##0',
                    'border': 1,
                    'align': 'right',
                    'valign': 'vcenter'
                })

                # Mapear subtotales a las columnas correctas
                for key, valor in subtotales.items():
                    # Extraer el nombre del campo del key
                    # Puede ser "SUM(presupuesto)" o directamente "Importe" (nuevo formato)
                    if '(' in key and ')' in key:
                        # Formato antiguo: "FUNCION(campo)"
                        campo_nombre = key.split('(')[1].rstrip(')')
                    else:
                        # Formato nuevo: nombre directo de columna
                        campo_nombre = key

                    # Determinar el formato según el tipo de agregación
                    formato_agg = formatos_agregaciones.get(key, 'ninguno')
                    formato_a_usar = formato_subtotal if formato_agg == 'moneda' else formato_subtotal_entero

                    # Buscar la columna correspondiente
                    if campo_nombre in columnas:
                        col_idx = columnas.index(campo_nombre)
                        worksheet.write(row, col_idx, valor, formato_a_usar)
                    elif campo_nombre == '*':
                        # COUNT(*) se escribe en la segunda columna, siempre como entero
                        worksheet.write(row, 1, valor, formato_subtotal_entero)

                row += 1

            row += 1  # Espacio entre grupos

        return row

    def _reemplazar_marcador(self, doc, marcador_texto: str, texto_reemplazo: str):
        """
        Reemplaza un marcador de TEXTO literal en el documento Word
        Maneja casos donde el marcador está dividido en múltiples runs

        Args:
            doc: Documento de Word
            marcador_texto: Texto del marcador a buscar (ej: "[TITULO_DEL_INFORME]")
            texto_reemplazo: Texto con el que reemplazar el marcador
        """
        reemplazado = False

        def reemplazar_en_paragrafo(paragraph):
            """Función helper para reemplazar en un párrafo"""
            if marcador_texto in paragraph.text:
                # El marcador está en este párrafo (puede estar dividido en múltiples runs)
                # Reconstruir todo el texto del párrafo
                texto_completo = paragraph.text
                nuevo_texto = texto_completo.replace(marcador_texto, texto_reemplazo)

                # Guardar el formato del primer run (si existe)
                primer_run_formato = None
                if paragraph.runs:
                    primer_run = paragraph.runs[0]
                    primer_run_formato = {
                        'bold': primer_run.bold,
                        'italic': primer_run.italic,
                        'font_name': primer_run.font.name,
                        'font_size': primer_run.font.size,
                        'font_color': primer_run.font.color.rgb if primer_run.font.color.rgb else None
                    }

                # Limpiar todos los runs del párrafo
                for run in paragraph.runs:
                    run.text = ""

                # Crear un nuevo run con el texto reemplazado
                if paragraph.runs:
                    # Usar el primer run existente
                    paragraph.runs[0].text = nuevo_texto
                else:
                    # Crear un nuevo run
                    nuevo_run = paragraph.add_run(nuevo_texto)
                    # Aplicar formato guardado si existe
                    if primer_run_formato:
                        if primer_run_formato['bold'] is not None:
                            nuevo_run.bold = primer_run_formato['bold']
                        if primer_run_formato['italic'] is not None:
                            nuevo_run.italic = primer_run_formato['italic']
                        if primer_run_formato['font_name']:
                            nuevo_run.font.name = primer_run_formato['font_name']
                        if primer_run_formato['font_size']:
                            nuevo_run.font.size = primer_run_formato['font_size']
                        if primer_run_formato['font_color']:
                            nuevo_run.font.color.rgb = primer_run_formato['font_color']

                return True
            return False

        # Buscar en párrafos del cuerpo
        for paragraph in doc.paragraphs:
            if reemplazar_en_paragrafo(paragraph):
                reemplazado = True

        # Buscar en encabezados
        for idx_section, section in enumerate(doc.sections):
            # Buscar en párrafos del encabezado
            for idx_para, paragraph in enumerate(section.header.paragraphs):
                if reemplazar_en_paragrafo(paragraph):
                    reemplazado = True

            # Buscar también en tablas del encabezado (importante!)
            for idx_table, table in enumerate(section.header.tables):
                for idx_row, row in enumerate(table.rows):
                    for idx_cell, cell in enumerate(row.cells):
                        for idx_para, paragraph in enumerate(cell.paragraphs):
                            if reemplazar_en_paragrafo(paragraph):
                                reemplazado = True

            # Buscar en pies de página
            for paragraph in section.footer.paragraphs:
                if reemplazar_en_paragrafo(paragraph):
                    reemplazado = True

        if not reemplazado:
            print(f"Advertencia: No se encontró el marcador '{marcador_texto}'")

    def _insertar_tabla_en_marcador(
        self,
        doc,
        marcador_texto: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        orientacion: str = 'horizontal',
        tipo_informe: str = ''
    ):
        """
        Inserta una tabla en la posición del marcador de TEXTO literal

        Args:
            doc: Documento de Word
            marcador_texto: Texto del marcador donde insertar la tabla (ej: "[Tabla_de_datos]")
            columnas: Lista de nombres de columnas
            datos: Datos a insertar
            resultado_agrupacion: Estructura de agrupaciones (opcional)
            orientacion: Orientación de la página ('horizontal' o 'vertical')
            tipo_informe: Tipo de informe para anchos personalizados
        """
        # Buscar el párrafo que contiene el marcador de texto
        target_paragraph = None

        # Buscar en párrafos del cuerpo
        for paragraph in doc.paragraphs:
            if marcador_texto in paragraph.text:
                target_paragraph = paragraph
                break

        # Si no se encuentra en párrafos, buscar en tablas
        if target_paragraph is None:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if marcador_texto in paragraph.text:
                                target_paragraph = paragraph
                                break
                        if target_paragraph:
                            break
                    if target_paragraph:
                        break
                if target_paragraph:
                    break

        if target_paragraph is None:
            print(f"Advertencia: No se encontró el marcador '{marcador_texto}'")
            return

        # Si hay grupos, procesar con estructura de agrupaciones
        if resultado_agrupacion and resultado_agrupacion.get('grupos'):
            # Remover el párrafo del marcador
            p_element = target_paragraph._element
            p_element.getparent().remove(p_element)

            # Exportar grupos directamente en el documento
            modo = resultado_agrupacion.get('modo', 'detalle')
            self._exportar_grupos_word(doc, resultado_agrupacion['grupos'], columnas, modo, 0, resultado_agrupacion, orientacion, tipo_informe)

            # Totales generales
            if resultado_agrupacion.get('totales_generales'):
                p_total = doc.add_paragraph()
                run_total = p_total.add_run("═══ TOTAL EJECUCIÓN MATERIAL ═══")
                run_total.font.bold = True
                run_total.font.size = Pt(12)

                # Mostrar totales
                formatos_agregaciones = resultado_agrupacion.get('formatos_agregaciones', {})
                totales = resultado_agrupacion['totales_generales']

                for key, valor in totales.items():
                    if '(' in key and ')' in key:
                        campo_nombre = key.split('(')[1].rstrip(')')
                    else:
                        campo_nombre = key

                    formato_agg = formatos_agregaciones.get(key, 'ninguno')
                    if formato_agg == 'moneda':
                        valor_texto = f"{valor:,.2f} €"
                    else:
                        valor_texto = f"{valor:,.2f}" if isinstance(valor, (int, float)) else str(valor)

                    run_total.add_text(f"  {campo_nombre}={valor_texto}")

            return

        # Insertar tabla después del marcador (sin agrupaciones)
        # Calcular número de filas necesarias
        num_filas = 1 + len(datos)  # Encabezado + datos

        # Crear tabla (se añade al final del documento)
        table = doc.add_table(rows=num_filas, cols=len(columnas))
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.allow_autofit = False

        # Calcular ancho disponible según orientación
        # A4 horizontal = 29.7cm - márgenes 1.5cm×2 = 26.7cm
        # A4 vertical = 21cm - márgenes 1.5cm×2 = 18cm
        ancho_disponible = 18.0 if orientacion == 'vertical' else 26.7

        # Ajustar ancho de tabla al ancho de página
        table.width = Cm(ancho_disponible)
        self._set_table_width(table, ancho_disponible)

        # Anchos personalizados para informes de Recursos
        anchos_recursos = {
            'Código': 1.5,
            'Cantidad': 2.0,
            'Ud.': 1.0,
            'Recurso / Material': 9.5,
            'Precio unitario': 2.0,
            'Importe': 2.0
        }

        # Determinar si usar anchos personalizados
        usa_anchos_personalizados = (
            'Recursos' in tipo_informe and
            all(col in anchos_recursos for col in columnas)
        )

        # Calcular anchos de columnas
        if usa_anchos_personalizados:
            # Usar anchos personalizados para informes de Recursos
            col_widths = [anchos_recursos[col] for col in columnas]
        else:
            # Distribución equitativa
            col_width_cm = ancho_disponible / len(columnas)
            col_widths = [col_width_cm] * len(columnas)

        # Establecer ancho usando XML directo
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                self._set_cell_width(cell, col_widths[col_idx])

        # Encabezados
        header_cells = table.rows[0].cells
        for idx, col_name in enumerate(columnas):
            cell = header_cells[idx]
            cell.text = col_name
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Tahoma'
                    run.font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Datos
        formatos_columnas = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}
        for fila_idx, fila_datos in enumerate(datos):
            row_cells = table.rows[fila_idx + 1].cells
            for col_idx, valor in enumerate(fila_datos):
                cell = row_cells[col_idx]

                # Obtener el formato del campo según su columna
                col_name = columnas[col_idx] if col_idx < len(columnas) else None
                formato_campo = formatos_columnas.get(col_name, 'ninguno') if col_name else 'ninguno'

                # Detectar si es una coordenada (latitud/longitud)
                es_coordenada = col_name and ('latitud' in col_name.lower() or 'longitud' in col_name.lower())

                # Aplicar formato según el tipo de campo
                if isinstance(valor, (int, float, Decimal)):
                    # Convertir Decimal a float para poder formatear
                    if isinstance(valor, Decimal):
                        valor = float(valor)

                    if es_coordenada:
                        # Coordenadas geográficas: 4 decimales
                        cell.text = f"{valor:.4f}"
                    elif formato_campo == 'moneda':
                        cell.text = f"{valor:,.2f} €"
                    elif formato_campo == 'decimal':
                        cell.text = f"{valor:,.2f}"
                    else:
                        # Por defecto, números con 2 decimales
                        cell.text = f"{valor:,.2f}"
                else:
                    cell.text = str(valor) if valor is not None else ""

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(8)

                    # Alinear a la derecha los campos numéricos
                    if col_name and col_name in ['Cantidad', 'Precio unitario', 'Importe']:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # MOVER la tabla a la posición del marcador
        tbl_element = table._element
        p_element = target_paragraph._element

        # Insertar la tabla ANTES del párrafo que contiene el marcador
        p_element.addprevious(tbl_element)

        # Eliminar completamente el párrafo del marcador (sustituirlo por la tabla)
        p_element.getparent().remove(p_element)

    def exportar_a_word(
        self,
        filepath: str,
        informe_nombre: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        proyecto_nombre: str = "",
        proyecto_codigo: str = "",
        fecha_informe: str = "",
        tipo_informe: Optional[str] = None
    ) -> bool:
        """
        Exporta el informe a Word usando plantilla profesional

        Args:
            filepath: Ruta del archivo Word a crear
            informe_nombre: Nombre del informe
            columnas: Lista de nombres de columnas
            datos: Datos del informe
            resultado_agrupacion: Estructura de agrupaciones y totales (opcional)
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto
            fecha_informe: Fecha del informe
            tipo_informe: Tipo de informe para seleccionar plantilla específica

        Returns:
            True si la exportación fue exitosa
        """
        try:
            # Importar configuración de plantillas y PDF
            from script.plantillas_config import obtener_ruta_plantilla
            from script.pdf_config import obtener_configuracion_pdf
            from script.informes_config import INFORMES_DEFINICIONES

            # Obtener configuración del informe
            config_pdf = obtener_configuracion_pdf(tipo_informe or informe_nombre)
            definicion = INFORMES_DEFINICIONES.get(tipo_informe or informe_nombre, {})
            campos_fijos = definicion.get('campos_fijos', False)

            # Si hay columnas_datos en resultado_agrupacion, usarlas (sin columna de agrupación)
            columnas_para_word = columnas
            datos_para_word = datos
            if resultado_agrupacion and resultado_agrupacion.get('columnas_datos'):
                columnas_para_word = resultado_agrupacion['columnas_datos']

                # Filtrar datos para que coincidan con columnas_datos
                # Crear mapeo de índices: columnas_datos -> columnas originales
                indices_columnas = []
                for col in columnas_para_word:
                    if col in columnas:
                        indices_columnas.append(columnas.index(col))

                # Filtrar cada fila de datos usando los índices
                datos_para_word = [
                    tuple(fila[idx] for idx in indices_columnas)
                    for fila in datos
                ]

            # Obtener plantilla apropiada según el tipo de informe
            # Si no se especifica tipo, usar el nombre del informe
            plantilla_path = obtener_ruta_plantilla(tipo_informe or informe_nombre)

            # Verificar si existe la plantilla
            if not os.path.exists(plantilla_path):
                print(f"Advertencia: No se encontró la plantilla en {plantilla_path}")
                print("Usando creación manual de documento...")
                # Fallback: crear documento vacío
                doc = Document()
                # TODO: Implementar creación manual si no hay plantilla
            else:
                # IMPORTANTE: Copiar la plantilla primero, no abrirla directamente
                # para no modificar el archivo original
                import shutil

                # Copiar plantilla al destino final
                shutil.copy2(plantilla_path, filepath)

                # Ahora abrir la copia para modificarla
                doc = Document(filepath)

            # Configurar orientación de página según config
            orientacion = config_pdf.get('orientacion', 'horizontal')
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENT

            for section in doc.sections:
                if orientacion == 'vertical':
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Inches(8.27)   # 21 cm
                    section.page_height = Inches(11.69)  # 29.7 cm
                    # Ajustar ancho de tablas en encabezado para vertical
                    ancho_disponible = 18.0  # 21cm - 3cm márgenes
                else:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.69)
                    section.page_height = Inches(8.27)
                    # Ajustar ancho de tablas en encabezado para horizontal
                    ancho_disponible = 26.7  # 29.7cm - 3cm márgenes

                # Ajustar ancho de las tablas del encabezado
                for table in section.header.tables:
                    self._set_table_width(table, ancho_disponible)

            # Reemplazar marcadores en la plantilla
            # Usar fecha proporcionada por el usuario o generar automáticamente
            fecha_actual = fecha_informe if fecha_informe else datetime.now().strftime("%d/%m/%Y")

            # Reemplazar marcadores de texto literal
            self._reemplazar_marcador(doc, "[TITULO_DEL_INFORME]", informe_nombre.upper())
            self._reemplazar_marcador(doc, "[FECHA]", fecha_actual)

            # Insertar tabla en el marcador (con o sin agrupaciones)
            self._insertar_tabla_en_marcador(
                doc,
                "[TABLA_DE_DATOS]",
                columnas_para_word,  # Usar columnas filtradas
                datos_para_word,     # Usar datos filtrados
                resultado_agrupacion,
                orientacion,  # Pasar orientación
                tipo_informe or informe_nombre  # Pasar tipo para anchos personalizados
            )

            # Guardar el documento
            doc.save(filepath)
            return True

        except Exception as e:
            print(f"Error al exportar a Word: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _exportar_grupos_word(
        self,
        doc: Document,
        grupos: List[Dict],
        columnas: List[str],
        modo: str = 'detalle',
        nivel: int = 0,
        resultado_agrupacion: Optional[Dict] = None,
        orientacion: str = 'horizontal',
        tipo_informe: str = ''
    ):
        """Exporta grupos jerárquicos a Word (recursivo)"""
        for grupo in grupos:
            clave = grupo.get('clave', '')
            campo = grupo.get('campo', '')
            datos = grupo.get('datos', [])
            subtotales = grupo.get('subtotales', {})
            subgrupos = grupo.get('subgrupos')
            nivel_grupo = grupo.get('nivel', nivel)

            # Encabezado del grupo
            indent = "    " * nivel_grupo
            titulo_grupo = f"{indent}📁 {campo.upper()}: {clave}"

            p_grupo = doc.add_paragraph(titulo_grupo)
            run_grupo = p_grupo.runs[0]
            run_grupo.font.name = 'Calibri'
            run_grupo.font.bold = True

            # Color y tamaño según nivel
            if nivel_grupo == 0:
                run_grupo.font.size = Pt(12)
                run_grupo.font.color.rgb = RGBColor(74, 111, 165)
            elif nivel_grupo == 1:
                run_grupo.font.size = Pt(11)
                run_grupo.font.color.rgb = RGBColor(107, 143, 184)
            else:
                run_grupo.font.size = Pt(10)
                run_grupo.font.color.rgb = RGBColor(138, 173, 199)

            # Si hay subgrupos, procesarlos recursivamente
            if subgrupos:
                self._exportar_grupos_word(doc, subgrupos, columnas, modo, nivel_grupo, resultado_agrupacion, orientacion, tipo_informe)
            elif modo == 'detalle' and datos:
                # Crear tabla para los datos
                table = doc.add_table(rows=1 + len(datos), cols=len(columnas))
                table.style = 'Light List Accent 1'
                table.autofit = False
                table.allow_autofit = False

                # Calcular ancho disponible según orientación
                # A4 horizontal = 29.7cm - márgenes 1.5cm×2 = 26.7cm
                # A4 vertical = 21cm - márgenes 1.5cm×2 = 18cm
                ancho_disponible = 18.0 if orientacion == 'vertical' else 26.7

                # Ajustar ancho de tabla al ancho de página
                table.width = Cm(ancho_disponible)
                # Forzar ancho de tabla usando XML directo
                self._set_table_width(table, ancho_disponible)

                # Anchos personalizados para informes de Recursos
                anchos_recursos = {
                    'Código': 1.5,
                    'Cantidad': 2.0,
                    'Ud.': 1.0,
                    'Recurso / Material': 9.5,
                    'Precio unitario': 2.0,
                    'Importe': 2.0
                }

                # Determinar si usar anchos personalizados
                usa_anchos_personalizados = (
                    'Recursos' in tipo_informe and
                    all(col in anchos_recursos for col in columnas)
                )

                # Calcular anchos de columnas
                if usa_anchos_personalizados:
                    # Usar anchos personalizados para informes de Recursos
                    col_widths = [anchos_recursos[col] for col in columnas]
                else:
                    # Distribución equitativa
                    col_width_cm = ancho_disponible / len(columnas)
                    col_widths = [col_width_cm] * len(columnas)

                # Encabezados
                header_cells = table.rows[0].cells
                for idx, col_name in enumerate(columnas):
                    cell = header_cells[idx]
                    cell.text = col_name
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Tahoma'
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Establecer ancho de celda directamente en XML
                    self._set_cell_width(cell, col_widths[idx])

                # Datos
                formatos_columnas = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}
                for row_idx, fila_datos in enumerate(datos, start=1):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, valor in enumerate(fila_datos):
                        cell = row_cells[col_idx]

                        # Obtener el formato del campo según su columna
                        col_name = columnas[col_idx] if col_idx < len(columnas) else None
                        formato_campo = formatos_columnas.get(col_name, 'ninguno') if col_name else 'ninguno'

                        # Detectar si es una coordenada (latitud/longitud)
                        es_coordenada = col_name and ('latitud' in col_name.lower() or 'longitud' in col_name.lower())

                        # Aplicar formato según el tipo de campo
                        if isinstance(valor, (int, float)):
                            if es_coordenada:
                                # Coordenadas geográficas: 4 decimales
                                cell.text = f"{valor:.4f}"
                            elif formato_campo == 'moneda':
                                cell.text = f"{valor:,.2f} €"
                            else:
                                # Por defecto, números con 2 decimales
                                cell.text = f"{valor:,.2f}"
                        else:
                            cell.text = str(valor) if valor is not None else ""

                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Tahoma'
                                run.font.size = Pt(8)

                            # Alinear a la derecha los campos numéricos
                            if col_name and col_name in ['Cantidad', 'Precio unitario', 'Importe']:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        # Establecer ancho de celda directamente en XML
                        self._set_cell_width(cell, col_widths[col_idx])

            # Subtotales
            if subtotales:
                indent_subtotal = "    " * (nivel_grupo + 1)
                p_subtotal = doc.add_paragraph(f"{indent_subtotal}▸ Subtotal")
                run_subtotal = p_subtotal.runs[0]
                run_subtotal.font.name = 'Tahoma'
                run_subtotal.font.size = Pt(9)
                run_subtotal.font.bold = True

                for key, valor in subtotales.items():
                    valor_texto = f"{valor:,.2f} €" if isinstance(valor, (int, float)) else str(valor)
                    run_subtotal.add_text(f"  {key}={valor_texto}")

            doc.add_paragraph()  # Espacio entre grupos

    def _set_cell_background(self, cell, color_hex: str):
        """Establece el color de fondo de una celda en Word"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_width(self, cell, width_cm: float):
        """
        Establece el ancho de una celda en Word directamente en el XML

        Args:
            cell: Celda de la tabla de Word
            width_cm: Ancho deseado en centímetros
        """
        # Convertir cm a twips (twentieths of a point)
        # 1 cm = 567 twips
        width_twips = int(width_cm * 567)

        # Obtener o crear el elemento tcPr (table cell properties)
        tcPr = cell._element.get_or_add_tcPr()

        # Buscar si ya existe un elemento tcW
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            # Crear nuevo elemento tcW
            tcW = OxmlElement('w:tcW')
            tcPr.append(tcW)

        # Establecer el ancho
        tcW.set(qn('w:w'), str(width_twips))
        tcW.set(qn('w:type'), 'dxa')  # dxa = twips (formato estándar)

    def _set_table_width(self, table, width_cm: float):
        """
        Establece el ancho de una tabla en Word directamente en el XML

        Args:
            table: Tabla de Word
            width_cm: Ancho deseado en centímetros
        """
        # Convertir cm a twips (twentieths of a point)
        # 1 cm = 567 twips
        width_twips = int(width_cm * 567)

        # Obtener el elemento tbl
        tbl = table._element

        # Obtener o crear tblPr (table properties)
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Buscar si ya existe un elemento tblW
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            # Crear nuevo elemento tblW
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)

        # Establecer el ancho
        tblW.set(qn('w:w'), str(width_twips))
        tblW.set(qn('w:type'), 'dxa')  # dxa = twips (formato estándar)

        # Asegurar que autofit esté desactivado
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')  # fixed layout, no autofit

    def _add_page_number(self, run):
        """Añade número de página a Word"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def exportar_a_pdf(
        self,
        filepath: str,
        informe_nombre: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        proyecto_nombre: str = "",
        proyecto_codigo: str = "",
        fecha_informe: str = "",
        tipo_informe: Optional[str] = None
    ) -> bool:
        """
        Exporta el informe a PDF usando ReportLab con control total del diseño

        Args:
            filepath: Ruta del archivo PDF a crear
            informe_nombre: Nombre del informe
            columnas: Lista de nombres de columnas
            datos: Datos del informe
            resultado_agrupacion: Estructura de agrupaciones y totales (opcional)
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto
            fecha_informe: Fecha del informe
            tipo_informe: Tipo de informe para seleccionar plantilla específica

        Returns:
            True si la exportación fue exitosa
        """
        try:
            from script.pdf_agrupaciones import PDFAgrupaciones
            from script.pdf_config import obtener_configuracion_pdf, aplicar_configuracion_a_plantilla
            from script.informes_config import INFORMES_DEFINICIONES

            print(f"Generando PDF con ReportLab: {filepath}")

            # Obtener configuración específica del tipo de informe
            config = obtener_configuracion_pdf(tipo_informe or informe_nombre)

            # Obtener definición del informe para verificar campos_fijos
            definicion = INFORMES_DEFINICIONES.get(tipo_informe or informe_nombre, {})
            campos_fijos = definicion.get('campos_fijos', False)
            usar_agregacion_sql = definicion.get('usar_agregacion_sql', False)
            campos_def = definicion.get('campos', {})

            # Crear plantilla PDF con la configuración
            pdf = PDFAgrupaciones(
                schema=self.schema,
                orientacion=config.get('orientacion', 'horizontal'),
                titulo=informe_nombre,
                proyecto_nombre=proyecto_nombre,
                proyecto_codigo=proyecto_codigo,
                fecha=fecha_informe
            )

            # Aplicar configuración de colores y estilos
            pdf = aplicar_configuracion_a_plantilla(pdf, config)

            # DETECTAR SI ES INFORME ESPECIAL DE ÓRDENES CON RECURSOS
            es_ordenes_recursos = (resultado_agrupacion and
                                   resultado_agrupacion.get('tipo') == 'ordenes_con_recursos')

            if es_ordenes_recursos:
                # Renderizado especial para órdenes con recursos
                print("DEBUG: Renderizando PDF especial de Órdenes con Recursos")
                return self._exportar_pdf_ordenes_recursos(
                    filepath, resultado_agrupacion, config, proyecto_nombre,
                    proyecto_codigo, fecha_informe, informe_nombre
                )

            # Agregar encabezado con logos (si está configurado)
            if config.get('mostrar_logos', True):
                pdf.agregar_encabezado()

            # Agregar información del proyecto (si está configurado)
            if config.get('mostrar_proyecto', True):
                pdf.agregar_info_proyecto(mostrar_fecha=config.get('mostrar_fecha', True))

            # Si el informe tiene campos_fijos, necesitamos filtrar las columnas
            # para NO mostrar campos de agrupación que se añaden solo para GROUP BY
            columnas_filtradas = list(columnas)
            datos_filtrados = datos
            formatos_filtrados = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}

            # Si hay grupos, usar columnas_datos que ya viene calculado desde informes.py
            if resultado_agrupacion and resultado_agrupacion.get('grupos'):
                # Usar columnas_datos si existe (grupos con SQL agregation)
                if 'columnas_datos' in resultado_agrupacion:
                    columnas_filtradas = resultado_agrupacion['columnas_datos']
                    # Los datos dentro de grupos ya están filtrados
                    # No necesitamos filtrar datos_filtrados porque se usan los datos de cada grupo
                    formatos_filtrados = {col: formatos_filtrados.get(col, 'ninguno') for col in columnas_filtradas}
                    print(f"DEBUG: Usando columnas_datos de resultado_agrupacion: {columnas_filtradas}")
                # Fallback al método anterior si no existe columnas_datos
                elif campos_fijos and resultado_agrupacion.get('agrupaciones'):
                    # Identificar columnas de agrupación por su nombre
                    columnas_a_eliminar = []
                    for agrup in resultado_agrupacion.get('agrupaciones', []):
                        # Buscar el nombre de la columna de agrupación
                        campo_def = campos_def.get(agrup, {})
                        nombre_columna = campo_def.get('nombre', agrup)
                        if nombre_columna in columnas_filtradas:
                            columnas_a_eliminar.append(nombre_columna)

                    # Filtrar columnas y datos
                    if columnas_a_eliminar:
                        indices_a_mantener = [i for i, col in enumerate(columnas) if col not in columnas_a_eliminar]
                        columnas_filtradas = [columnas[i] for i in indices_a_mantener]
                        datos_filtrados = [tuple(fila[i] for i in indices_a_mantener) for fila in datos]
                        formatos_filtrados = {col: formatos_filtrados.get(col, 'ninguno') for col in columnas_filtradas}

                        print(f"DEBUG: Filtradas {len(columnas_a_eliminar)} columnas de agrupación: {columnas_a_eliminar}")
                        print(f"DEBUG: Columnas finales para PDF: {columnas_filtradas}")

            # Agregar tabla de datos
            if resultado_agrupacion and resultado_agrupacion.get('grupos'):
                # Tabla con agrupaciones
                modo = resultado_agrupacion.get('modo', 'detalle')
                elementos_tabla = pdf.crear_tabla_agrupada(
                    columnas=columnas_filtradas,
                    datos=datos_filtrados,
                    resultado_agrupacion=resultado_agrupacion,
                    modo=modo
                )
                pdf.elements.extend(elementos_tabla)
            else:
                # Tabla simple sin agrupaciones
                tabla = pdf.crear_tabla_simple(
                    columnas=columnas_filtradas,
                    datos=datos_filtrados,
                    formatos_columnas=formatos_filtrados
                )
                pdf.elements.append(tabla)

                # Para informes de Recursos, añadir totales finales con GG y BI
                if usar_agregacion_sql and 'Importe' in columnas_filtradas:
                    # Calcular total de ejecución material (columna Importe)
                    idx_importe = columnas_filtradas.index('Importe')
                    total_ejecucion_material = 0.0
                    for fila in datos_filtrados:
                        valor = fila[idx_importe]
                        if valor is not None:
                            from decimal import Decimal
                            total_ejecucion_material += float(valor) if isinstance(valor, Decimal) else valor

                    # Añadir espacio antes de totales
                    from reportlab.platypus import Spacer
                    from reportlab.lib.units import cm
                    pdf.elements.append(Spacer(1, 0.5 * cm))

                    # Crear y añadir tabla de totales finales
                    tabla_totales = pdf.crear_tabla_totales_finales(
                        total_ejecucion_material=total_ejecucion_material,
                        columnas=columnas_filtradas
                    )
                    pdf.elements.append(tabla_totales)

            # Agregar pie de página
            pdf.agregar_pie_pagina()

            # Generar PDF
            exito = pdf.generar_pdf(filepath)

            if exito:
                print(f"✓ PDF generado correctamente: {filepath}")
                return True
            else:
                print("✗ Error al generar PDF")
                return False

        except Exception as e:
            print(f"Error al exportar a PDF: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _exportar_pdf_ordenes_recursos(
        self,
        filepath: str,
        resultado_ordenes: Dict,
        config: Dict,
        proyecto_nombre: str,
        proyecto_codigo: str,
        fecha_informe: str,
        informe_nombre: str
    ) -> bool:
        """
        Renderiza PDF especial para informe de Órdenes de Trabajo con Recursos

        Formato:
        - Encabezado con logos + título
        - Para cada orden:
            * Cabecera con datos de la orden (código, título, fecha, localización, coordenadas)
            * Tabla de recursos presupuestados
            * Total de la orden
        - Si hay agrupaciones, mostrar subtotales por grupo
        - Gran total al final
        - Pie de página con fecha y paginación

        Args:
            filepath: Ruta del PDF a generar
            resultado_ordenes: Estructura con órdenes y recursos
            config: Configuración PDF del tipo de informe
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto
            fecha_informe: Fecha del informe
            informe_nombre: Nombre del informe

        Returns:
            True si se generó correctamente
        """
        try:
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib import colors as reportlab_colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from datetime import datetime

            print(f"Generando PDF de Órdenes con Recursos: {filepath}")

            # Crear documento
            # topMargin: 0.5cm desde borde + 2cm encabezado + 0.5cm separación = 3.0cm
            doc = SimpleDocTemplate(
                filepath,
                pagesize=A4,  # Vertical (Portrait)
                rightMargin=1.5*cm,
                leftMargin=1.5*cm,
                topMargin=3.0*cm,
                bottomMargin=2*cm
            )

            # Estilos (tonos de grises)
            styles = getSampleStyleSheet()
            style_titulo = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=reportlab_colors.HexColor('#404040'),  # Gris oscuro
                spaceAfter=12,
                alignment=TA_CENTER
            )
            style_orden_header = ParagraphStyle(
                'OrdenHeader',
                parent=styles['Normal'],
                fontSize=8,
                textColor=reportlab_colors.HexColor('#606060'),  # Gris medio
                spaceBefore=10,
                spaceAfter=6,
                fontName='Helvetica-Bold'
            )
            style_normal = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=7,
                spaceAfter=4
            )
            style_grupo = ParagraphStyle(
                'GrupoHeader',
                parent=styles['Heading2'],
                fontSize=10,
                textColor=reportlab_colors.white,
                backColor=reportlab_colors.HexColor('#606060'),  # Gris medio
                spaceBefore=12,
                spaceAfter=8,
                fontName='Helvetica-Bold'
            )

            # Elementos del PDF
            elements = []

            # Obtener datos
            ordenes = resultado_ordenes.get('ordenes', [])
            grupos = resultado_ordenes.get('grupos', [])
            agrupaciones = resultado_ordenes.get('agrupaciones', [])
            gran_total = resultado_ordenes.get('gran_total', 0)

            # Función auxiliar para renderizar una orden
            def renderizar_orden(orden_data):
                """Renderiza una orden con su cabecera y tabla de recursos"""
                elementos_orden = []

                datos_orden = orden_data['datos_orden']
                recursos = orden_data['recursos']
                total_orden = orden_data['total_orden']

                # CABECERA DE LA ORDEN
                # Primera fila: Código + Localización (NO en tabla, como párrafo)
                codigo = datos_orden.get('codigo', '')
                titulo = datos_orden.get('titulo', '')

                # Crear párrafo para código y localización
                style_codigo = ParagraphStyle(
                    'CodigoStyle',
                    parent=style_normal,
                    fontSize=8,
                    fontName='Helvetica-Bold',
                    spaceAfter=6
                )
                codigo_titulo_text = f"<b>{codigo}</b>   {titulo}"
                elementos_orden.append(Paragraph(codigo_titulo_text, style_codigo))
                elementos_orden.append(Spacer(1, 0.3*cm))

                # Resto de campos de la orden (FECHA, LOCALIZACIÓN, LATITUD/LONGITUD)
                fecha_fin = datos_orden.get('fecha_fin', '')
                if isinstance(fecha_fin, datetime):
                    fecha_fin = fecha_fin.strftime('%d/%m/%Y')
                elif isinstance(fecha_fin, date):
                    fecha_fin = fecha_fin.strftime('%d/%m/%Y')

                municipio = datos_orden.get('municipio', '')
                localizacion = datos_orden.get('localizacion', '')
                latitud = datos_orden.get('latitud', '')
                longitud = datos_orden.get('longitud', '')

                # Formatear latitud y longitud a 4 decimales
                if latitud:
                    try:
                        latitud = f"{float(latitud):.4f}"
                    except (ValueError, TypeError):
                        latitud = str(latitud)
                if longitud:
                    try:
                        longitud = f"{float(longitud):.4f}"
                    except (ValueError, TypeError):
                        longitud = str(longitud)

                # Combinar municipio + localización
                loc_completa = f"{municipio}"
                if localizacion:
                    loc_completa += f" - {localizacion}"

                # Tabla con el resto de datos
                # Fila 2: Fecha (3.5cm + 14.5cm)
                # Fila 3: Localización (3.5cm + 14.5cm)
                # Fila 4: Latitud y Longitud (3.5cm + 5.5cm + 3.5cm + 5.5cm = 18cm)
                datos_cabecera = [
                    [Paragraph("<b>Fecha:</b>", style_normal), Paragraph(str(fecha_fin), style_normal)],
                    [Paragraph("<b>Localización:</b>", style_normal), Paragraph(loc_completa, style_normal)],
                ]

                # Anchos para filas 2 y 3 (2 columnas)
                col_widths_2cols = [3.5*cm, 14.5*cm]

                # Crear tabla para Fecha y Localización
                tabla_fecha_loc = Table(datos_cabecera, colWidths=col_widths_2cols)
                tabla_fecha_loc.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BACKGROUND', (0, 0), (-1, -1), reportlab_colors.HexColor('#E8E8E8')),
                    ('BOX', (0, 0), (-1, -1), 1, reportlab_colors.HexColor('#808080')),
                    ('INNERGRID', (0, 0), (-1, -1), 0.5, reportlab_colors.HexColor('#808080')),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                elementos_orden.append(tabla_fecha_loc)
                elementos_orden.append(Spacer(1, 0.3*cm))

                # Fila 4: Latitud y Longitud (4 columnas: 3.5cm + 5.5cm + 3.5cm + 5.5cm)
                if latitud and longitud:
                    datos_coordenadas = [[
                        Paragraph("<b>Latitud:</b>", style_normal),
                        Paragraph(str(latitud), style_normal),
                        Paragraph("<b>Longitud:</b>", style_normal),
                        Paragraph(str(longitud), style_normal)
                    ]]
                    col_widths_4cols = [3.5*cm, 5.5*cm, 3.5*cm, 5.5*cm]

                    tabla_coordenadas = Table(datos_coordenadas, colWidths=col_widths_4cols)
                    tabla_coordenadas.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BACKGROUND', (0, 0), (-1, -1), reportlab_colors.HexColor('#E8E8E8')),
                        ('BOX', (0, 0), (-1, -1), 1, reportlab_colors.HexColor('#808080')),
                        ('INNERGRID', (0, 0), (-1, -1), 0.5, reportlab_colors.HexColor('#808080')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ]))
                    elementos_orden.append(tabla_coordenadas)
                    elementos_orden.append(Spacer(1, 0.3*cm))

                # Más espacio antes de la tabla de recursos
                elementos_orden.append(Spacer(1, 0.4*cm))

                # TABLA DE RECURSOS
                if recursos:
                    # Encabezados
                    headers = ['Código', 'Cantidad', 'Ud.', 'Recurso / Material', 'Precio unit.', 'Importe']
                    tabla_datos = [headers]

                    # Filas de recursos
                    for recurso in recursos:
                        # Usar Paragraph para Recurso/Material para permitir multilínea
                        recurso_text = Paragraph(recurso.get('resumen', ''), style_normal)
                        fila = [
                            recurso.get('codigo', ''),
                            f"{recurso.get('cantidad', 0):.2f}",
                            recurso.get('unidad', ''),
                            recurso_text,
                            f"{recurso.get('coste', 0):.2f} €",
                            f"{recurso.get('coste_total', 0):.2f} €"
                        ]
                        tabla_datos.append(fila)

                    # Fila de total - "Total Parte" con span
                    tabla_datos.append([
                        Paragraph(f"<b>Total Parte {codigo}</b>", style_normal), '', '', '', '',
                        Paragraph(f"<b>{total_orden:.2f} €</b>", style_normal)
                    ])

                    # Anchos de columnas (total 18cm)
                    col_widths = [1.5*cm, 2*cm, 1*cm, 9.5*cm, 2*cm, 2*cm]

                    tabla_recursos = Table(tabla_datos, colWidths=col_widths)
                    tabla_recursos.setStyle(TableStyle([
                        # Encabezado
                        ('BACKGROUND', (0, 0), (-1, 0), reportlab_colors.HexColor('#A0A0A0')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), reportlab_colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 7),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),

                        # Datos
                        ('ALIGN', (0, 1), (2, -2), 'CENTER'),
                        ('ALIGN', (3, 1), (3, -2), 'LEFT'),
                        ('ALIGN', (4, 1), (5, -2), 'RIGHT'),
                        ('FONTSIZE', (0, 1), (-1, -2), 6),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [reportlab_colors.white, reportlab_colors.HexColor('#F5F5F5')]),
                        ('VALIGN', (0, 1), (-1, -2), 'TOP'),

                        # Fila de total - span desde columna 0 hasta 4
                        ('SPAN', (0, -1), (4, -1)),
                        ('BACKGROUND', (0, -1), (-1, -1), reportlab_colors.HexColor('#D0D0D0')),
                        ('ALIGN', (0, -1), (0, -1), 'LEFT'),
                        ('ALIGN', (5, -1), (5, -1), 'RIGHT'),
                        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),

                        # Bordes
                        ('BOX', (0, 0), (-1, -1), 1, reportlab_colors.black),
                        ('INNERGRID', (0, 0), (-1, -2), 0.5, reportlab_colors.grey),

                        # Padding
                        ('LEFTPADDING', (0, 0), (-1, -1), 4),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                        ('TOPPADDING', (0, 0), (-1, -1), 3),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ]))

                    elementos_orden.append(tabla_recursos)
                else:
                    # No hay recursos
                    elementos_orden.append(Paragraph("<i>Sin recursos presupuestados</i>", style_normal))

                # Más espacio entre las partes (órdenes)
                elementos_orden.append(Spacer(1, 1*cm))
                return elementos_orden

            # Función auxiliar para renderizar grupos recursivamente
            def renderizar_grupos(grupos_lista, nivel=0):
                """Renderiza grupos jerárquicamente"""
                elementos_grupos = []

                for grupo in grupos_lista:
                    campo = grupo['campo']
                    valor = grupo['valor']
                    subtotal = grupo['subtotal']
                    ordenes_grupo = grupo.get('ordenes', [])
                    subgrupos = grupo.get('subgrupos', [])

                    # Color según nivel (tonos de grises)
                    if nivel == 0:
                        color_fondo = config.get('color_grupo_nivel0', '#505050')
                    elif nivel == 1:
                        color_fondo = config.get('color_grupo_nivel1', '#707070')
                    else:
                        color_fondo = config.get('color_grupo_nivel2', '#909090')

                    # Header del grupo
                    style_grupo_nivel = ParagraphStyle(
                        f'Grupo_Nivel_{nivel}',
                        parent=style_grupo,
                        backColor=reportlab_colors.HexColor(color_fondo)
                    )
                    elementos_grupos.append(Paragraph(f"<b>{campo.upper()}: {valor}</b>", style_grupo_nivel))
                    elementos_grupos.append(Spacer(1, 0.2*cm))

                    # Si hay subgrupos, renderizarlos
                    if subgrupos:
                        elementos_grupos.extend(renderizar_grupos(subgrupos, nivel + 1))
                    # Si no hay subgrupos, renderizar las órdenes
                    elif ordenes_grupo:
                        for orden in ordenes_grupo:
                            elementos_grupos.extend(renderizar_orden(orden))

                    # Subtotal del grupo
                    tabla_subtotal = Table(
                        [[Paragraph(f"<b>SUBTOTAL {campo.upper()}: {valor}</b>", style_normal),
                          Paragraph(f"<b>{subtotal:.2f} €</b>", style_normal)]],
                        colWidths=[None, 3*cm]
                    )
                    tabla_subtotal.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), reportlab_colors.HexColor('#D0D0D0')),
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                        ('BOX', (0, 0), (-1, -1), 1, reportlab_colors.black),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ]))
                    elementos_grupos.append(tabla_subtotal)
                    elementos_grupos.append(Spacer(1, 0.5*cm))

                return elementos_grupos

            # Renderizar contenido
            if grupos:
                # Con agrupaciones
                elements.extend(renderizar_grupos(grupos))
            else:
                # Sin agrupaciones, mostrar órdenes directamente
                for orden in ordenes:
                    elements.extend(renderizar_orden(orden))

            # GRAN TOTAL
            if config.get('mostrar_gran_total', True):
                tabla_gran_total = Table(
                    [[Paragraph("<b>TOTAL GENERAL:</b>", style_titulo),
                      Paragraph(f"<b>{gran_total:.2f} €</b>", style_titulo)]],
                    colWidths=[None, 4*cm]
                )
                tabla_gran_total.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), reportlab_colors.HexColor('#505050')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), reportlab_colors.white),
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 12),
                    ('BOX', (0, 0), (-1, -1), 2, reportlab_colors.black),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ]))
                elements.append(Spacer(1, 0.5*cm))
                elements.append(tabla_gran_total)

            # ENCABEZADO (en todas las páginas)
            def encabezado_pagina(canvas, doc):
                """Función para agregar encabezado con logos y título en cada página"""
                canvas.saveState()

                # Altura del encabezado: 2cm
                # Margen desde el borde: 0.5cm
                y_encabezado = A4[1] - 0.5*cm  # Posición Y desde el borde superior

                # Dibujar logos con altura de 2cm (mantener aspect ratio)
                logo_height = 2*cm
                logo_y = y_encabezado - logo_height

                # Logo izquierdo
                if self.logo_redes_path and os.path.exists(self.logo_redes_path):
                    try:
                        # Obtener dimensiones originales de la imagen
                        img = PILImage.open(self.logo_redes_path)
                        img_width, img_height = img.size
                        aspect_ratio = img_width / img_height

                        # Calcular ancho manteniendo aspect ratio
                        logo_width = logo_height * aspect_ratio

                        # Dibujar logo izquierdo
                        canvas.drawImage(self.logo_redes_path,
                                       1.5*cm, logo_y,
                                       width=logo_width, height=logo_height,
                                       preserveAspectRatio=True, mask='auto')
                    except Exception as e:
                        print(f"Error al cargar logo izquierdo: {e}")

                # Logo derecho
                if self.logo_urbide_path and os.path.exists(self.logo_urbide_path):
                    try:
                        # Obtener dimensiones originales de la imagen
                        img = PILImage.open(self.logo_urbide_path)
                        img_width, img_height = img.size
                        aspect_ratio = img_width / img_height

                        # Calcular ancho manteniendo aspect ratio
                        logo_width = logo_height * aspect_ratio

                        # Dibujar logo derecho (alineado a la derecha)
                        canvas.drawImage(self.logo_urbide_path,
                                       A4[0] - 1.5*cm - logo_width, logo_y,
                                       width=logo_width, height=logo_height,
                                       preserveAspectRatio=True, mask='auto')
                    except Exception as e:
                        print(f"Error al cargar logo derecho: {e}")

                # Título centrado
                canvas.setFont('Helvetica-Bold', 18)
                canvas.setFillColor(reportlab_colors.HexColor('#404040'))
                titulo_y = y_encabezado - logo_height / 2
                canvas.drawCentredString(A4[0] / 2, titulo_y, informe_nombre.upper())

                canvas.restoreState()

            # PIE DE PÁGINA (con fecha y paginación)
            def pie_pagina(canvas, doc):
                """Función para agregar pie de página en cada página"""
                canvas.saveState()
                fecha_str = fecha_informe or datetime.now().strftime('%d/%m/%Y')
                pagina_str = f"Página {doc.page} de {doc.page}"  # Se actualizará en el segundo paso

                # Fecha a la izquierda
                canvas.setFont('Helvetica', 9)
                canvas.drawString(2*cm, 1.5*cm, fecha_str)

                # Paginación a la derecha
                canvas.drawRightString(A4[0] - 2*cm, 1.5*cm, f"Página {canvas.getPageNumber()}")
                canvas.restoreState()

            # Construir PDF
            doc.build(elements, onFirstPage=lambda c, d: (encabezado_pagina(c, d), pie_pagina(c, d)),
                     onLaterPages=lambda c, d: (encabezado_pagina(c, d), pie_pagina(c, d)))

            print(f"✓ PDF de Órdenes con Recursos generado: {filepath}")
            return True

        except Exception as e:
            print(f"Error al generar PDF de Órdenes con Recursos: {e}")
            import traceback
            traceback.print_exc()
            return False

    def exportar_a_pdf_word(
        self,
        filepath: str,
        informe_nombre: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        proyecto_nombre: str = "",
        proyecto_codigo: str = "",
        fecha_informe: str = "",
        tipo_informe: Optional[str] = None
    ) -> bool:
        """
        [LEGACY] Exporta el informe a PDF generando primero un Word y convirtiéndolo a PDF
        Esta función se mantiene por compatibilidad pero ya no se usa por defecto

        Args:
            filepath: Ruta del archivo PDF a crear
            informe_nombre: Nombre del informe
            columnas: Lista de nombres de columnas
            datos: Datos del informe
            resultado_agrupacion: Estructura de agrupaciones y totales (opcional)
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto
            fecha_informe: Fecha del informe
            tipo_informe: Tipo de informe para seleccionar plantilla específica

        Returns:
            True si la exportación fue exitosa
        """
        import tempfile

        try:
            # Paso 1: Generar el archivo Word temporal
            print("Generando archivo Word temporal...")
            temp_word = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_word_path = temp_word.name
            temp_word.close()

            # Usar la misma función de exportar_a_word (MISMO .docx para Word y PDF)
            exito_word = self.exportar_a_word(
                filepath=temp_word_path,
                informe_nombre=informe_nombre,
                columnas=columnas,
                datos=datos,
                resultado_agrupacion=resultado_agrupacion,
                proyecto_nombre=proyecto_nombre,
                proyecto_codigo=proyecto_codigo,
                fecha_informe=fecha_informe,
                tipo_informe=tipo_informe
            )

            if not exito_word:
                print("Error al generar el archivo Word temporal")
                return False

            # Paso 2: Convertir Word a PDF usando el sistema
            print(f"Convirtiendo Word a PDF: {temp_word_path} -> {filepath}")

            # Intentar diferentes métodos de conversión
            exito_conversion = self._convertir_word_a_pdf(temp_word_path, filepath)

            # Paso 3: Limpiar archivo temporal
            try:
                os.unlink(temp_word_path)
            except:
                pass

            if exito_conversion:
                print(f"✓ PDF generado correctamente: {filepath}")
                return True
            else:
                print("✗ Error al convertir Word a PDF")
                return False

        except Exception as e:
            print(f"Error al exportar a PDF: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _convertir_word_a_pdf(self, word_path: str, pdf_path: str) -> bool:
        """
        Convierte un archivo Word a PDF usando diferentes métodos disponibles

        Args:
            word_path: Ruta del archivo Word de entrada
            pdf_path: Ruta del archivo PDF de salida

        Returns:
            True si la conversión fue exitosa
        """
        import platform

        # Método 1: Intentar con win32com (Microsoft Word COM)
        if platform.system() == 'Windows':
            try:
                import win32com.client
                import pythoncom

                pythoncom.CoInitialize()
                word = win32com.client.Dispatch('Word.Application')
                word.Visible = False

                # Abrir documento
                doc = word.Documents.Open(os.path.abspath(word_path))

                # Guardar como PDF (wdFormatPDF = 17)
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                doc.Close()
                word.Quit()

                pythoncom.CoUninitialize()

                print("✓ Conversión exitosa usando Microsoft Word COM")
                return True

            except Exception as e:
                print(f"⚠ No se pudo usar Word COM: {e}")

        # Método 2: Intentar con LibreOffice (si está instalado)
        try:
            # Buscar LibreOffice en ubicaciones comunes
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "/usr/bin/libreoffice",
                "/usr/local/bin/libreoffice",
            ]

            soffice_path = None
            for path in libreoffice_paths:
                if os.path.exists(path):
                    soffice_path = path
                    break

            if soffice_path:
                cmd = [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(os.path.abspath(pdf_path)),
                    os.path.abspath(word_path)
                ]

                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                # LibreOffice genera el PDF con el mismo nombre que el Word
                generated_pdf = os.path.join(
                    os.path.dirname(os.path.abspath(pdf_path)),
                    os.path.splitext(os.path.basename(word_path))[0] + '.pdf'
                )

                if os.path.exists(generated_pdf):
                    # Mover al destino final si es diferente
                    if generated_pdf != os.path.abspath(pdf_path):
                        import shutil
                        shutil.move(generated_pdf, pdf_path)

                    print("✓ Conversión exitosa usando LibreOffice")
                    return True

        except Exception as e:
            print(f"⚠ No se pudo usar LibreOffice: {e}")

        # Método 3: Mensaje de error si ningún método funcionó
        print("\n" + "="*70)
        print("ERROR: No se pudo convertir el documento Word a PDF")
        print("="*70)
        print("\nSoluciones posibles:")
        print("1. Instalar Microsoft Word")
        print("2. Instalar LibreOffice: https://www.libreoffice.org/download/")
        print("3. Instalar la librería win32com: pip install pywin32")
        print("\nEl archivo Word se guardó correctamente en:")
        print(f"   {word_path}")
        print("\nPuede abrir este archivo y guardarlo manualmente como PDF.")
        print("="*70 + "\n")

        return False

    def exportar_a_pdf_old(
        self,
        filepath: str,
        informe_nombre: str,
        columnas: List[str],
        datos: List[tuple],
        resultado_agrupacion: Optional[Dict] = None,
        proyecto_nombre: str = "",
        proyecto_codigo: str = ""
    ) -> bool:
        """
        [VERSIÓN ANTIGUA] Exporta el informe a PDF usando ReportLab
        Esta función se mantiene por compatibilidad pero ya no se usa

        Args:
            filepath: Ruta del archivo PDF a crear
            informe_nombre: Nombre del informe
            columnas: Lista de nombres de columnas
            datos: Datos del informe
            resultado_agrupacion: Estructura de agrupaciones y totales (opcional)
            proyecto_nombre: Nombre del proyecto
            proyecto_codigo: Código del proyecto

        Returns:
            True si la exportación fue exitosa
        """
        try:
            # Crear el documento PDF en orientación horizontal
            # A4 landscape = 29.7cm x 21cm
            # Márgenes: 1.5cm a cada lado
            # Ancho disponible = 29.7cm - 1.5cm - 1.5cm = 26.7cm
            doc = SimpleDocTemplate(
                filepath,
                pagesize=landscape(A4),
                topMargin=1.5*cm,
                bottomMargin=1.5*cm,
                leftMargin=1.5*cm,
                rightMargin=1.5*cm
            )

            # Lista de elementos del documento
            elements = []

            # Estilos
            styles = getSampleStyleSheet()

            # Calcular ancho disponible para tablas
            ancho_pagina = landscape(A4)[0]  # 29.7cm en puntos
            ancho_disponible = ancho_pagina - (1.5*cm * 2)  # Restar márgenes (26.7cm en puntos)

            # Estilo para título
            style_titulo = ParagraphStyle(
                'Titulo',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.black,
                alignment=TA_CENTER,
                spaceAfter=12
            )

            # Estilo para subtítulo
            style_subtitulo = ParagraphStyle(
                'Subtitulo',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#7C7C7C'),
                italic=True,
                alignment=TA_LEFT,
                spaceAfter=6
            )

            # Crear tabla de encabezado con logos y título
            header_data = []
            header_row = []

            # Logo izquierdo (Logo Redes Urbide) - 2.1cm de altura
            if self.logo_redes_path and os.path.exists(self.logo_redes_path):
                img_left = Image(self.logo_redes_path, height=2.1*cm, width=None)
                header_row.append(img_left)
            else:
                header_row.append('')

            # Título en el centro
            titulo_para = Paragraph(informe_nombre.upper(), style_titulo)
            header_row.append(titulo_para)

            # Logo derecho (Logo Urbide) - 2.1cm de altura
            if self.logo_urbide_path and os.path.exists(self.logo_urbide_path):
                img_right = Image(self.logo_urbide_path, height=2.1*cm, width=None)
                header_row.append(img_right)
            else:
                header_row.append('')

            header_data.append(header_row)

            # Tabla de encabezado con anchos: 3.5cm, resto, 3.5cm
            # Ancho disponible = 26.7cm (calculado arriba)
            ancho_titulo = ancho_disponible - (3.5*cm * 2)  # 26.7cm - 7cm = 19.7cm
            header_table = Table(header_data, colWidths=[3.5*cm, ancho_titulo, 3.5*cm])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),   # Logo izquierdo alineado a la izquierda
                ('ALIGN', (1, 0), (1, 0), 'CENTER'), # Título centrado
                ('ALIGN', (2, 0), (2, 0), 'RIGHT'),  # Logo derecho alineado a la derecha
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(header_table)
            elements.append(Spacer(1, 0.5*cm))

            # Información del proyecto
            if proyecto_nombre:
                p_proyecto = Paragraph(proyecto_nombre, style_subtitulo)
                elements.append(p_proyecto)

            # Fecha
            fecha_actual = datetime.now().strftime("%d/%m/%Y")
            style_fecha = ParagraphStyle(
                'Fecha',
                parent=styles['Normal'],
                fontSize=10,
                fontName='Helvetica-Bold',
                spaceAfter=12
            )
            p_fecha = Paragraph(f"FECHA: {fecha_actual}", style_fecha)
            elements.append(p_fecha)
            elements.append(Spacer(1, 0.5*cm))

            # Preparar datos de la tabla
            if resultado_agrupacion and resultado_agrupacion.get('grupos'):
                # Si hay agrupaciones, crear estructura jerárquica
                tabla_datos = self._crear_tabla_grupos_pdf(
                    resultado_agrupacion['grupos'],
                    columnas,
                    resultado_agrupacion.get('modo', 'detalle'),
                    resultado_agrupacion,
                    ancho_disponible
                )
                elements.extend(tabla_datos)

                # Totales generales
                if resultado_agrupacion.get('totales_generales'):
                    elements.append(Spacer(1, 0.3*cm))

                    style_total = ParagraphStyle(
                        'Total',
                        parent=styles['Normal'],
                        fontSize=12,
                        fontName='Helvetica-Bold',
                        spaceAfter=6
                    )
                    p_total = Paragraph("═══ TOTAL GENERAL ═══", style_total)
                    elements.append(p_total)

                    totales = resultado_agrupacion['totales_generales']
                    for key, valor in totales.items():
                        if isinstance(valor, (int, float)):
                            texto = f"{key}: {valor:,.2f} €"
                        else:
                            texto = f"{key}: {valor}"
                        p = Paragraph(texto, styles['Normal'])
                        elements.append(p)
            else:
                # Tabla simple sin agrupaciones
                tabla_datos = [columnas]

                # Formatear datos
                formatos_columnas = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}
                for fila in datos:
                    fila_formateada = []
                    for col_idx, valor in enumerate(fila):
                        col_name = columnas[col_idx] if col_idx < len(columnas) else None
                        formato_campo = formatos_columnas.get(col_name, 'ninguno') if col_name else 'ninguno'

                        if isinstance(valor, (int, float)):
                            if formato_campo == 'moneda':
                                fila_formateada.append(f"{valor:,.2f} €")
                            else:
                                fila_formateada.append(f"{valor:,.2f}")
                        else:
                            fila_formateada.append(str(valor) if valor is not None else "")
                    tabla_datos.append(fila_formateada)

                # Calcular anchos de columnas dinámicos
                num_columnas = len(columnas)
                ancho_por_columna = ancho_disponible / num_columnas
                col_widths = [ancho_por_columna] * num_columnas

                # Crear tabla con anchos dinámicos
                tabla = Table(tabla_datos, colWidths=col_widths)
                tabla.setStyle(TableStyle([
                    # Encabezado
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9D9D9')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    # Datos
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    # Bordes
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(tabla)

            # Pie de página
            elements.append(Spacer(1, 1*cm))
            style_footer = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#7C7C7C'),
                italic=True,
                alignment=TA_CENTER
            )
            p_footer = Paragraph(
                f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}",
                style_footer
            )
            elements.append(p_footer)

            # Construir el PDF
            doc.build(elements)
            print(f"✓ PDF generado correctamente: {filepath}")
            return True

        except Exception as e:
            print(f"Error al exportar a PDF: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _crear_tabla_grupos_pdf(
        self,
        grupos: List[Dict],
        columnas: List[str],
        modo: str = 'detalle',
        resultado_agrupacion: Optional[Dict] = None,
        ancho_disponible: float = 24*cm
    ) -> List:
        """Crea elementos de tabla para grupos jerárquicos en PDF"""
        elements = []
        styles = getSampleStyleSheet()

        for grupo in grupos:
            clave = grupo.get('clave', '')
            campo = grupo.get('campo', '')
            datos = grupo.get('datos', [])
            subtotales = grupo.get('subtotales', {})
            subgrupos = grupo.get('subgrupos')
            nivel = grupo.get('nivel', 0)

            # Estilo del grupo según nivel
            if nivel == 0:
                bg_color = colors.HexColor('#4A6FA5')
                font_size = 12
            elif nivel == 1:
                bg_color = colors.HexColor('#6B8FB8')
                font_size = 11
            else:
                bg_color = colors.HexColor('#8AADC7')
                font_size = 10

            # Título del grupo
            indent = "    " * nivel
            titulo_grupo = f"{indent}📁 {campo.upper()}: {clave}"

            style_grupo = ParagraphStyle(
                f'Grupo{nivel}',
                parent=styles['Normal'],
                fontSize=font_size,
                fontName='Helvetica-Bold',
                textColor=colors.white,
                spaceAfter=6
            )

            # Crear tabla para el título del grupo
            grupo_data = [[Paragraph(titulo_grupo, style_grupo)]]
            grupo_table = Table(grupo_data, colWidths=[ancho_disponible])
            grupo_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), bg_color),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), font_size),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(grupo_table)
            elements.append(Spacer(1, 0.2*cm))

            # Si hay subgrupos, procesarlos recursivamente
            if subgrupos:
                sub_elements = self._crear_tabla_grupos_pdf(
                    subgrupos,
                    columnas,
                    modo,
                    resultado_agrupacion,
                    ancho_disponible
                )
                elements.extend(sub_elements)
            elif modo == 'detalle' and datos:
                # Crear tabla de datos
                tabla_datos = [columnas]
                formatos_columnas = resultado_agrupacion.get('formatos_columnas', {}) if resultado_agrupacion else {}

                for fila in datos:
                    fila_formateada = []
                    for col_idx, valor in enumerate(fila):
                        col_name = columnas[col_idx] if col_idx < len(columnas) else None
                        formato_campo = formatos_columnas.get(col_name, 'ninguno') if col_name else 'ninguno'

                        if isinstance(valor, (int, float)):
                            if formato_campo == 'moneda':
                                fila_formateada.append(f"{valor:,.2f} €")
                            else:
                                fila_formateada.append(f"{valor:,.2f}")
                        else:
                            fila_formateada.append(str(valor) if valor is not None else "")
                    tabla_datos.append(fila_formateada)

                # Calcular anchos de columnas dinámicos
                num_columnas = len(columnas)
                ancho_por_columna = ancho_disponible / num_columnas
                col_widths = [ancho_por_columna] * num_columnas

                tabla = Table(tabla_datos, colWidths=col_widths)
                tabla.setStyle(TableStyle([
                    # Encabezado
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9D9D9')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    # Datos
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    # Bordes
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(tabla)
                elements.append(Spacer(1, 0.2*cm))

            # Subtotales
            if subtotales:
                indent_subtotal = "    " * (nivel + 1)
                texto_subtotales = f"{indent_subtotal}▸ Subtotal"

                for key, valor in subtotales.items():
                    if isinstance(valor, (int, float)):
                        texto_subtotales += f"  {key}={valor:,.2f} €"
                    else:
                        texto_subtotales += f"  {key}={valor}"

                style_subtotal = ParagraphStyle(
                    'Subtotal',
                    parent=styles['Normal'],
                    fontSize=9,
                    fontName='Helvetica-Bold',
                    spaceAfter=6
                )

                subtotal_data = [[Paragraph(texto_subtotales, style_subtotal)]]
                subtotal_table = Table(subtotal_data, colWidths=[ancho_disponible])
                subtotal_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E7E6E6')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(subtotal_table)

            elements.append(Spacer(1, 0.3*cm))

        return elements
