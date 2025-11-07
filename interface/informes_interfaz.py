# interface/informes_interfaz.py
"""
Interfaz del Módulo de Informes
Generación de informes personalizados con filtros multicriterio
"""

import customtkinter
from tkinter import ttk, filedialog
from tkcalendar import DateEntry
from PIL import Image
import os
import sys
import datetime

# Agregar el directorio padre al path para imports
current_path = os.path.dirname(os.path.realpath(__file__))
parent_path = os.path.dirname(current_path)
sys.path.append(parent_path)

from script.informes_config import (
    CATEGORIAS_INFORMES,
    INFORMES_DEFINICIONES,
    CAMPOS_PARTES,
    CAMPOS_RECURSOS,
    CAMPOS_PRESUPUESTOS,
    CAMPOS_CERTIFICACIONES,
    CAMPOS_PLANIFICACION,
    OPERADORES,
    TIPOS_CAMPO,
    FORMATOS_SALIDA,
    ORDEN_OPCIONES,
    LOGICA_FILTROS,
    CONFIG_CABECERA_DEFAULT
)
from script.informes import get_dimension_values, ejecutar_informe, ejecutar_informe_con_agrupacion
from script.informes_storage import InformesConfigStorage


class InformesFrame(customtkinter.CTkFrame):
    """Frame principal del módulo de Informes"""

    def __init__(self, master, user, password, schema, informe_inicial=None, **kwargs):
        super().__init__(master, **kwargs)

        self.user = user
        self.password = password
        self.schema = schema
        self.informe_inicial = informe_inicial  # Informe a preseleccionar

        # Variables de estado
        self.informe_seleccionado = None
        self.categoria_seleccionada = None
        self.definicion_actual = None  # Definición completa del informe seleccionado
        self.ordenaciones = []
        self.filtros = []
        self.campos_seleccionados = {}
        self.agrupaciones = []  # Lista de campos por los que agrupar
        self.agregaciones = []  # Lista de agregaciones configuradas
        self.modo_visualizacion = "detalle"  # "detalle" o "resumen"

        # Gestor de almacenamiento de configuraciones
        self.storage = InformesConfigStorage()

        # Configurar grid - Header compacto + contenido principal + action bar
        self.grid_columnconfigure(0, weight=0)  # Panel izquierdo fijo
        self.grid_columnconfigure(1, weight=1)  # Panel derecho expandible
        self.grid_rowconfigure(0, weight=0)     # Header compacto (altura fija ~40px)
        self.grid_rowconfigure(1, weight=1)     # Row principal (contenido expandible)
        self.grid_rowconfigure(2, weight=0)     # Action bar (altura fija ~50px)

        # Crear componentes
        self._create_compact_header()
        self._create_left_panel()
        self._create_right_panel()
        self._create_action_bar()

        # Seleccionar informe inicial si se proporcionó
        if self.informe_inicial:
            self.after(100, lambda: self._select_initial_report(self.informe_inicial))

    def _create_compact_header(self):
        """Crea el header compacto en una sola línea con título y botón configuración"""
        # Frame del header (altura fija ~40px)
        header_frame = customtkinter.CTkFrame(self, fg_color="transparent", height=40)
        header_frame.grid(row=0, column=0, columnspan=2, sticky="ew", padx=15, pady=(8, 5))
        header_frame.grid_columnconfigure(0, weight=1)
        header_frame.grid_propagate(False)  # Mantener altura fija

        # Título a la izquierda
        title = customtkinter.CTkLabel(
            header_frame,
            text="GENERACIÓN DE INFORMES",
            font=customtkinter.CTkFont(size=16, weight="bold"),
            anchor="w"
        )
        title.grid(row=0, column=0, sticky="w")

        # Botón configuración a la derecha
        config_button = customtkinter.CTkButton(
            header_frame,
            text="⚙️ Configuración",
            width=120,
            height=30,
            font=customtkinter.CTkFont(size=11),
            command=self._open_config_dialog
        )
        config_button.grid(row=0, column=1, sticky="e")

        # Separador visual debajo del header
        separator = customtkinter.CTkFrame(self, height=1, fg_color="gray30")
        separator.grid(row=0, column=0, columnspan=2, sticky="ews", padx=15, pady=(38, 0))

    def _create_left_panel(self):
        """Crea el panel izquierdo con el árbol de informes - OCUPA TODA LA ALTURA"""
        left_frame = customtkinter.CTkFrame(self, width=300)
        left_frame.grid(row=1, column=0, sticky="nsew", padx=(15, 8), pady=(5, 5))  # row=1 (contenido)
        left_frame.grid_propagate(False)
        left_frame.grid_rowconfigure(1, weight=1)  # TreeView expande verticalmente

        # Título
        title = customtkinter.CTkLabel(
            left_frame,
            text="TIPO DE INFORME",
            font=customtkinter.CTkFont(size=13, weight="bold")
        )
        title.grid(row=0, column=0, padx=10, pady=(8, 5), sticky="w")

        # Frame para el TreeView
        tree_frame = customtkinter.CTkFrame(left_frame)
        tree_frame.grid(row=1, column=0, padx=10, pady=(5, 10), sticky="nsew")

        # TreeView con scroll
        self.tree_scroll = customtkinter.CTkScrollbar(tree_frame)
        self.tree_scroll.pack(side="right", fill="y")

        # Crear TreeView
        self.tree = ttk.Treeview(
            tree_frame,
            yscrollcommand=self.tree_scroll.set,
            show="tree"
        )
        self.tree.pack(side="left", fill="both", expand=True)
        self.tree_scroll.configure(command=self.tree.yview)

        # IMPORTANTE: Configurar el estilo DESPUÉS de crear el widget y empaquetarlo
        # Esto asegura que sobrescribe cualquier configuración global previa
        self._configure_tree_style()

        # Poblar el árbol
        self._populate_tree()

        # Bind para selección
        self.tree.bind('<<TreeviewSelect>>', self._on_tree_select)

    def _configure_tree_style(self):
        """Configura el estilo del TreeView DESPUÉS de crearlo para asegurar que se aplique"""
        import tkinter.font as tkfont

        # Crear estilo único DESPUÉS de que el TreeView existe
        style = ttk.Style()

        # Nombre único para evitar conflictos
        style_name = "InformesCustom.Treeview"

        # Configurar con fuente grande
        style.configure(style_name,
                        background="#2a2d2e",
                        foreground="white",
                        fieldbackground="#2a2d2e",
                        borderwidth=0,
                        rowheight=32,  # Fila más alta
                        font=('Segoe UI', 14, 'bold'))  # Fuente AÚN MÁS GRANDE para que sea VISIBLE

        style.configure(f"{style_name}.Heading",
                        background="#1f538d",
                        foreground="white",
                        font=('Segoe UI', 14, 'bold'))

        style.map(style_name,
                  background=[('selected', '#1f538d')],
                  foreground=[('selected', 'white')])

        # Aplicar el estilo al TreeView
        self.tree.configure(style=style_name)

        # ENFOQUE ALTERNATIVO: Configurar fuente directamente usando tags
        # Esto es más robusto que depender solo de Style
        custom_font = tkfont.Font(family='Segoe UI', size=14, weight='bold')

        # Crear un tag con la fuente personalizada
        self.tree.tag_configure('custom_font', font=custom_font)

    def _populate_tree(self):
        """Puebla el TreeView con categorías e informes"""
        for categoria, informes in CATEGORIAS_INFORMES.items():
            # Insertar categoría con tag personalizado
            cat_id = self.tree.insert("", "end", text=categoria, open=False, tags=('custom_font',))

            # Insertar informes de la categoría con tag personalizado
            for informe in informes:
                self.tree.insert(cat_id, "end", text=f"  {informe}", tags=('custom_font',))

    def _on_tree_select(self, event):
        """Maneja la selección en el TreeView"""
        selection = self.tree.selection()
        if selection:
            item = selection[0]
            text = self.tree.item(item, "text")

            # Verificar si es categoría o informe
            parent = self.tree.parent(item)

            if parent:  # Es un informe
                # Obtener categoría padre
                parent_text = self.tree.item(parent, "text")
                self.categoria_seleccionada = parent_text
                self.informe_seleccionado = text.strip()

                # Cargar definición del informe (si existe)
                self.definicion_actual = INFORMES_DEFINICIONES.get(self.informe_seleccionado)

                # Actualizar título del informe
                if self.definicion_actual:
                    # Solo mostrar el nombre del informe, sin prefijo
                    titulo = self.informe_seleccionado
                    self.informe_title_label.configure(
                        text=titulo,
                        text_color="white"
                    )
                else:
                    self.informe_title_label.configure(
                        text=f"{self.informe_seleccionado} (En desarrollo)",
                        text_color="gray"
                    )

                # Limpiar filtros y ordenaciones anteriores
                self._clear_all_filtros()
                self._clear_all_ordenaciones()

                # Actualizar campos disponibles según definición del informe
                self._update_campos_disponibles()

    def _select_initial_report(self, informe_name):
        """Selecciona un informe específico en el tree"""
        # Buscar el informe en el tree
        for cat_id in self.tree.get_children():
            for informe_id in self.tree.get_children(cat_id):
                text = self.tree.item(informe_id, "text").strip()
                if text == informe_name:
                    # Expandir categoría padre
                    self.tree.item(cat_id, open=True)
                    # Seleccionar el informe
                    self.tree.selection_set(informe_id)
                    # Ver el elemento
                    self.tree.see(informe_id)
                    # Disparar el evento de selección
                    self._on_tree_select(None)
                    return
        print(f"Advertencia: Informe '{informe_name}' no encontrado en el árbol")

    def _create_right_panel(self):
        """Crea el panel derecho con la configuración del informe"""
        right_frame = customtkinter.CTkScrollableFrame(self)
        right_frame.grid(row=1, column=1, sticky="nsew", padx=(8, 15), pady=(5, 5))  # row=1 (contenido)
        right_frame.grid_columnconfigure(0, weight=1)

        # Título del informe seleccionado
        self.informe_title_label = customtkinter.CTkLabel(
            right_frame,
            text="Informe seleccionado: Ninguno",
            font=customtkinter.CTkFont(size=12, weight="bold"),
            text_color="gray",
            anchor="w"
        )
        self.informe_title_label.grid(row=0, column=0, padx=15, pady=(5, 8), sticky="w")

        # Separador
        separator1 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator1.grid(row=1, column=0, sticky="ew", padx=15, pady=(0, 8))

        # Sección CLASIFICACIÓN
        self._create_ordenacion_section(right_frame, row=2)

        # Separador
        separator2 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator2.grid(row=3, column=0, sticky="ew", padx=15, pady=10)

        # Sección AGRUPACIÓN (GROUP BY)
        self._create_agrupacion_section(right_frame, row=4)

        # Separador
        separator3 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator3.grid(row=5, column=0, sticky="ew", padx=15, pady=10)

        # Sección AGREGACIONES (funciones)
        self._create_agregaciones_section(right_frame, row=6)

        # Separador
        separator4 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator4.grid(row=7, column=0, sticky="ew", padx=15, pady=10)

        # Sección FILTROS
        self._create_filtros_section(right_frame, row=8)

        # Separador
        separator5 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator5.grid(row=9, column=0, sticky="ew", padx=15, pady=10)

        # Sección CAMPOS A MOSTRAR
        self._create_campos_section(right_frame, row=10)

        # Separador
        separator6 = customtkinter.CTkFrame(right_frame, height=2, fg_color="gray30")
        separator6.grid(row=11, column=0, sticky="ew", padx=15, pady=10)

        # Sección OPCIONES DE PRESENTACIÓN
        self._create_opciones_section(right_frame, row=12)

    def _create_ordenacion_section(self, parent, row):
        """Crea la sección de ordenación"""
        # Frame contenedor con altura mínima
        clasif_frame = customtkinter.CTkFrame(parent, fg_color="transparent", height=150)
        clasif_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        clasif_frame.grid_columnconfigure(0, weight=1)

        # Título
        title = customtkinter.CTkLabel(
            clasif_frame,
            text="📋 CLASIFICACIÓN (Agrupar por)",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 8))

        # Botón añadir
        add_button = customtkinter.CTkButton(
            clasif_frame,
            text="+ Añadir ordenación",
            width=150,
            height=28,
            command=self._add_ordenacion
        )
        add_button.grid(row=1, column=0, sticky="w", pady=(0, 8))

        # Frame para ordenaciones con scroll - ALTURA AUMENTADA
        self.ordenaciones_container = customtkinter.CTkScrollableFrame(
            clasif_frame,
            height=220,  # Aumentado de 100 a 220
            fg_color="transparent"
        )
        self.ordenaciones_container.grid(row=2, column=0, sticky="ew")

        self.ordenaciones_frame = customtkinter.CTkFrame(self.ordenaciones_container, fg_color="transparent")
        self.ordenaciones_frame.pack(fill="both", expand=True)
        self.ordenaciones_frame.grid_columnconfigure(0, weight=1)

    def _create_agrupacion_section(self, parent, row):
        """Crea la sección de agrupación (GROUP BY visual)"""
        # Frame contenedor
        agrup_frame = customtkinter.CTkFrame(parent, fg_color="transparent", height=150)
        agrup_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        agrup_frame.grid_columnconfigure(0, weight=1)

        # Título y descripción
        title = customtkinter.CTkLabel(
            agrup_frame,
            text="📊 AGRUPACIÓN (Organizar visualmente)",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 3))

        desc = customtkinter.CTkLabel(
            agrup_frame,
            text="Organiza los registros en grupos visuales. Los subtotales se calcularán por cada grupo.",
            font=customtkinter.CTkFont(size=10),
            text_color="gray60"
        )
        desc.grid(row=1, column=0, sticky="w", pady=(0, 8))

        # Frame horizontal para botón y selector de modo
        controls_frame = customtkinter.CTkFrame(agrup_frame, fg_color="transparent")
        controls_frame.grid(row=2, column=0, sticky="ew", pady=(0, 8))
        controls_frame.grid_columnconfigure(1, weight=1)

        # Botón añadir
        add_button = customtkinter.CTkButton(
            controls_frame,
            text="+ Añadir nivel de agrupación",
            width=180,
            height=28,
            command=self._add_agrupacion
        )
        add_button.grid(row=0, column=0, sticky="w", padx=(0, 10))

        # Selector de modo (Detalle vs Resumen)
        modo_frame = customtkinter.CTkFrame(controls_frame, fg_color="transparent")
        modo_frame.grid(row=0, column=1, sticky="e")

        modo_label = customtkinter.CTkLabel(
            modo_frame,
            text="Modo:",
            font=customtkinter.CTkFont(size=10)
        )
        modo_label.pack(side="left", padx=(0, 5))

        self.modo_selector = customtkinter.CTkSegmentedButton(
            modo_frame,
            values=["Detalle", "Resumen"],
            width=180,
            height=28,
            command=self._on_modo_changed
        )
        self.modo_selector.set("Detalle")
        self.modo_selector.pack(side="left")

        # Frame para agrupaciones con scroll
        self.agrupaciones_container = customtkinter.CTkScrollableFrame(
            agrup_frame,
            height=180,
            fg_color="transparent"
        )
        self.agrupaciones_container.grid(row=3, column=0, sticky="ew")

        self.agrupaciones_frame = customtkinter.CTkFrame(self.agrupaciones_container, fg_color="transparent")
        self.agrupaciones_frame.pack(fill="both", expand=True)
        self.agrupaciones_frame.grid_columnconfigure(0, weight=1)

    def _create_agregaciones_section(self, parent, row):
        """Crea la sección de agregaciones (funciones SUM, AVG, etc.)"""
        # Frame contenedor
        agreg_frame = customtkinter.CTkFrame(parent, fg_color="transparent", height=150)
        agreg_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        agreg_frame.grid_columnconfigure(0, weight=1)

        # Título y descripción
        title = customtkinter.CTkLabel(
            agreg_frame,
            text="🔢 AGREGACIONES (Funciones de cálculo)",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 3))

        desc = customtkinter.CTkLabel(
            agreg_frame,
            text="Calcula totales, promedios, etc. Se mostrarán como subtotales por grupo y total general.",
            font=customtkinter.CTkFont(size=10),
            text_color="gray60"
        )
        desc.grid(row=1, column=0, sticky="w", pady=(0, 8))

        # Botón añadir
        add_button = customtkinter.CTkButton(
            agreg_frame,
            text="+ Añadir agregación",
            width=150,
            height=28,
            command=self._add_agregacion
        )
        add_button.grid(row=2, column=0, sticky="w", pady=(0, 8))

        # Frame para agregaciones con scroll
        self.agregaciones_container = customtkinter.CTkScrollableFrame(
            agreg_frame,
            height=180,
            fg_color="transparent"
        )
        self.agregaciones_container.grid(row=3, column=0, sticky="ew")

        self.agregaciones_frame = customtkinter.CTkFrame(self.agregaciones_container, fg_color="transparent")
        self.agregaciones_frame.pack(fill="both", expand=True)
        self.agregaciones_frame.grid_columnconfigure(0, weight=1)

    def _create_filtros_section(self, parent, row):
        """Crea la sección de filtros"""
        # Frame contenedor con altura mínima
        filtros_frame = customtkinter.CTkFrame(parent, fg_color="transparent", height=150)
        filtros_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        filtros_frame.grid_columnconfigure(0, weight=1)

        # Título
        title = customtkinter.CTkLabel(
            filtros_frame,
            text="🔍 FILTROS (Mostrar solo si...)",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 8))

        # Botón añadir
        add_button = customtkinter.CTkButton(
            filtros_frame,
            text="+ Añadir filtro",
            width=150,
            height=28,
            command=self._add_filtro
        )
        add_button.grid(row=1, column=0, sticky="w", pady=(0, 8))

        # Frame para filtros con scroll - ALTURA AUMENTADA
        self.filtros_container = customtkinter.CTkScrollableFrame(
            filtros_frame,
            height=250,  # Aumentado de 120 a 250
            fg_color="transparent"
        )
        self.filtros_container.grid(row=2, column=0, sticky="ew")

        self.filtros_frame = customtkinter.CTkFrame(self.filtros_container, fg_color="transparent")
        self.filtros_frame.pack(fill="both", expand=True)
        self.filtros_frame.grid_columnconfigure(0, weight=1)

    def _create_campos_section(self, parent, row):
        """Crea la sección de campos a mostrar"""
        # Frame contenedor
        campos_frame = customtkinter.CTkFrame(parent, fg_color="transparent")
        campos_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        campos_frame.grid_columnconfigure(0, weight=1)

        # Título
        title = customtkinter.CTkLabel(
            campos_frame,
            text="📄 CAMPOS A MOSTRAR",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 8))

        # Frame scrollable para checkboxes
        self.campos_scrollable = customtkinter.CTkScrollableFrame(campos_frame, height=180)
        self.campos_scrollable.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        self.campos_scrollable.grid_columnconfigure(0, weight=1)
        self.campos_scrollable.grid_columnconfigure(1, weight=1)
        self.campos_scrollable.grid_columnconfigure(2, weight=1)

        # Mensaje inicial
        self.campos_message = customtkinter.CTkLabel(
            self.campos_scrollable,
            text="Seleccione un informe para ver los campos disponibles",
            text_color="gray"
        )
        self.campos_message.grid(row=0, column=0, columnspan=3, pady=20)

        # Botones de selección
        buttons_frame = customtkinter.CTkFrame(campos_frame, fg_color="transparent")
        buttons_frame.grid(row=2, column=0, sticky="w")

        select_all = customtkinter.CTkButton(
            buttons_frame,
            text="☑ Seleccionar todo",
            width=140,
            height=28,
            command=self._select_all_campos
        )
        select_all.grid(row=0, column=0, padx=(0, 10))

        deselect_all = customtkinter.CTkButton(
            buttons_frame,
            text="☐ Deseleccionar todo",
            width=140,
            height=28,
            command=self._deselect_all_campos
        )
        deselect_all.grid(row=0, column=1)

    def _create_opciones_section(self, parent, row):
        """Crea la sección de opciones de presentación"""
        # Frame contenedor
        opciones_frame = customtkinter.CTkFrame(parent, fg_color="transparent")
        opciones_frame.grid(row=row, column=0, sticky="ew", padx=15, pady=3)
        opciones_frame.grid_columnconfigure(0, weight=1)

        # Título
        title = customtkinter.CTkLabel(
            opciones_frame,
            text="🎨 OPCIONES DE PRESENTACIÓN",
            font=customtkinter.CTkFont(size=12, weight="bold")
        )
        title.grid(row=0, column=0, sticky="w", pady=(0, 8))

        # Formato de salida
        formato_label = customtkinter.CTkLabel(
            opciones_frame,
            text="Formato de salida:",
            font=customtkinter.CTkFont(size=12)
        )
        formato_label.grid(row=1, column=0, sticky="w", pady=(5, 5))

        self.formato_var = customtkinter.StringVar(value="Tabla")
        formato_frame = customtkinter.CTkFrame(opciones_frame, fg_color="transparent")
        formato_frame.grid(row=2, column=0, sticky="w", padx=(20, 0))

        for i, formato in enumerate(FORMATOS_SALIDA):
            radio = customtkinter.CTkRadioButton(
                formato_frame,
                text=formato,
                variable=self.formato_var,
                value=formato
            )
            radio.grid(row=0, column=i, padx=(0, 20))

        # Opciones de totales
        totales_label = customtkinter.CTkLabel(
            opciones_frame,
            text="Opciones de totales:",
            font=customtkinter.CTkFont(size=12)
        )
        totales_label.grid(row=3, column=0, sticky="w", pady=(15, 5))

        self.subtotales_var = customtkinter.BooleanVar(value=True)
        self.subtotales_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Mostrar subtotales por grupo",
            variable=self.subtotales_var
        )
        self.subtotales_check.grid(row=4, column=0, sticky="w", padx=(20, 0), pady=2)

        self.totales_var = customtkinter.BooleanVar(value=True)
        self.totales_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Mostrar totales generales",
            variable=self.totales_var
        )
        self.totales_check.grid(row=5, column=0, sticky="w", padx=(20, 0), pady=2)

        self.graficos_var = customtkinter.BooleanVar(value=False)
        self.graficos_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Incluir gráficos resumen",
            variable=self.graficos_var
        )
        self.graficos_check.grid(row=6, column=0, sticky="w", padx=(20, 0), pady=2)

        # Opciones adicionales
        adicionales_label = customtkinter.CTkLabel(
            opciones_frame,
            text="Opciones adicionales:",
            font=customtkinter.CTkFont(size=12)
        )
        adicionales_label.grid(row=7, column=0, sticky="w", pady=(15, 5))

        self.logo_var = customtkinter.BooleanVar(value=True)
        logo_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Incluir logo de empresa",
            variable=self.logo_var
        )
        logo_check.grid(row=8, column=0, sticky="w", padx=(20, 0), pady=2)

        self.fecha_var = customtkinter.BooleanVar(value=True)
        fecha_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Incluir fecha de generación",
            variable=self.fecha_var
        )
        fecha_check.grid(row=9, column=0, sticky="w", padx=(20, 0), pady=2)

        self.compacto_var = customtkinter.BooleanVar(value=False)
        compacto_check = customtkinter.CTkCheckBox(
            opciones_frame,
            text="Modo compacto (menos espaciado)",
            variable=self.compacto_var
        )
        compacto_check.grid(row=10, column=0, sticky="w", padx=(20, 0), pady=2)

    def _create_action_bar(self):
        """Crea la barra de acciones inferior"""
        action_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        action_frame.grid(row=2, column=0, columnspan=2, sticky="ew", padx=15, pady=(5, 10))  # row=2 (action bar)
        action_frame.grid_columnconfigure(0, weight=1)

        # Frame para centrar botones
        buttons_frame = customtkinter.CTkFrame(action_frame, fg_color="transparent")
        buttons_frame.grid(row=0, column=0)

        # Botones
        save_config_btn = customtkinter.CTkButton(
            buttons_frame,
            text="💾 Guardar Config",
            width=130,
            height=35,
            fg_color="#2B5797",
            hover_color="#1E3F6B",
            command=self._guardar_configuracion
        )
        save_config_btn.grid(row=0, column=0, padx=5)

        load_config_btn = customtkinter.CTkButton(
            buttons_frame,
            text="📂 Cargar Config",
            width=130,
            height=35,
            fg_color="#2B5797",
            hover_color="#1E3F6B",
            command=self._cargar_configuracion
        )
        load_config_btn.grid(row=0, column=1, padx=5)

        preview_btn = customtkinter.CTkButton(
            buttons_frame,
            text="👁️ Previsualizar",
            width=140,
            height=35,
            command=self._preview_report
        )
        preview_btn.grid(row=0, column=2, padx=5)

        word_btn = customtkinter.CTkButton(
            buttons_frame,
            text="📄 Word",
            width=120,
            height=35,
            command=self._export_word
        )
        word_btn.grid(row=0, column=3, padx=5)

        excel_btn = customtkinter.CTkButton(
            buttons_frame,
            text="📊 Excel",
            width=120,
            height=35,
            command=self._export_excel
        )
        excel_btn.grid(row=0, column=4, padx=5)

        pdf_btn = customtkinter.CTkButton(
            buttons_frame,
            text="📕 PDF",
            width=120,
            height=35,
            command=self._export_pdf
        )
        pdf_btn.grid(row=0, column=5, padx=5)

    def _add_ordenacion(self):
        """Añade un nuevo selector de ordenación dinámico"""
        if not self.definicion_actual:
            from CTkMessagebox import CTkMessagebox
            CTkMessagebox(
                title="Aviso",
                message="Selecciona primero un informe para poder clasificar los datos",
                icon="warning"
            )
            return

        row = len(self.ordenaciones)

        clasif_container = customtkinter.CTkFrame(self.ordenaciones_frame)
        clasif_container.grid(row=row, column=0, sticky="ew", pady=5)
        clasif_container.grid_columnconfigure(1, weight=1)
        clasif_container.grid_columnconfigure(3, weight=1)

        # Label
        label = customtkinter.CTkLabel(
            clasif_container,
            text=f"Ordenación {row + 1}:",
            font=customtkinter.CTkFont(size=11)
        )
        label.grid(row=0, column=0, padx=(0, 10), sticky="w")

        # Variable - Poblar con campos disponibles del informe
        var_label = customtkinter.CTkLabel(clasif_container, text="Variable:")
        var_label.grid(row=0, column=1, sticky="w", padx=(0, 5))

        # Obtener todos los campos del informe
        campos_informe = self.definicion_actual.get('campos', {})
        nombres_campos = [campo_def['nombre'] for campo_key, campo_def in campos_informe.items()]

        # Crear objeto de ordenación que actualizaremos
        clasif_obj = {
            'container': clasif_container,
            'var_combo': None,
            'orden_combo': None,
            'campo_actual': None  # Guardará el campo_key seleccionado
        }

        var_combo = customtkinter.CTkComboBox(
            clasif_container,
            values=nombres_campos if nombres_campos else ["Sin campos"],
            width=200,
            command=lambda choice: self._on_ordenacion_campo_change(clasif_obj, choice)
        )
        var_combo.grid(row=0, column=2, sticky="w", padx=(0, 20))
        clasif_obj['var_combo'] = var_combo

        # Orden
        orden_label = customtkinter.CTkLabel(clasif_container, text="Orden:")
        orden_label.grid(row=0, column=3, sticky="w", padx=(0, 5))

        orden_combo = customtkinter.CTkComboBox(
            clasif_container,
            values=ORDEN_OPCIONES,
            width=150
        )
        orden_combo.grid(row=0, column=4, sticky="w", padx=(0, 10))
        clasif_obj['orden_combo'] = orden_combo

        # Botón eliminar
        del_btn = customtkinter.CTkButton(
            clasif_container,
            text="🗑️",
            width=40,
            fg_color="darkred",
            hover_color="red",
            command=lambda: self._remove_ordenacion(clasif_container)
        )
        del_btn.grid(row=0, column=5, padx=(0, 5))

        # Añadir a la lista
        self.ordenaciones.append(clasif_obj)

        # Auto-seleccionar primer campo si hay campos disponibles
        if nombres_campos:
            var_combo.set(nombres_campos[0])
            self._on_ordenacion_campo_change(clasif_obj, nombres_campos[0])

    def _on_ordenacion_campo_change(self, clasif_obj, campo_nombre):
        """Maneja el cambio de campo en una ordenación"""
        if not self.definicion_actual:
            return

        campos_informe = self.definicion_actual.get('campos', {})

        # Buscar el campo_key que corresponde al nombre seleccionado
        for campo_key, campo_def in campos_informe.items():
            if campo_def['nombre'] == campo_nombre:
                clasif_obj['campo_actual'] = campo_key
                break

    def _remove_ordenacion(self, container):
        """Elimina un selector de ordenación"""
        container.destroy()
        # Actualizar lista
        self.ordenaciones = [c for c in self.ordenaciones if c['container'].winfo_exists()]

    def _add_agrupacion(self):
        """Añade un nuevo nivel de agrupación"""
        if not self.definicion_actual:
            from CTkMessagebox import CTkMessagebox
            CTkMessagebox(
                title="Aviso",
                message="Selecciona primero un informe para poder configurar agrupaciones",
                icon="warning"
            )
            return

        # Verificar límite de niveles
        agrupaciones_config = self.definicion_actual.get('agrupaciones', {})
        max_niveles = agrupaciones_config.get('max_niveles', 3)

        if len(self.agrupaciones) >= max_niveles:
            from CTkMessagebox import CTkMessagebox
            CTkMessagebox(
                title="Límite alcanzado",
                message=f"Máximo {max_niveles} niveles de agrupación permitidos",
                icon="warning"
            )
            return

        row = len(self.agrupaciones)

        agrup_container = customtkinter.CTkFrame(self.agrupaciones_frame, fg_color="gray25")
        agrup_container.grid(row=row, column=0, sticky="ew", pady=5, padx=5)
        agrup_container.grid_columnconfigure(1, weight=1)

        # Label de nivel
        nivel_label = customtkinter.CTkLabel(
            agrup_container,
            text=f"Nivel {row + 1}:",
            font=customtkinter.CTkFont(size=11, weight="bold")
        )
        nivel_label.grid(row=0, column=0, padx=(10, 10), sticky="w")

        # Campo para agrupar
        campo_label = customtkinter.CTkLabel(agrup_container, text="Agrupar por:")
        campo_label.grid(row=0, column=1, sticky="w", padx=(0, 5))

        # Obtener campos permitidos para agrupación
        campos_permitidos = agrupaciones_config.get('campos_permitidos', [])
        campos_informe = self.definicion_actual.get('campos', {})

        nombres_campos = [
            campos_informe[campo_key]['nombre']
            for campo_key in campos_permitidos
            if campo_key in campos_informe
        ]

        # Crear objeto de agrupación
        agrup_obj = {
            'container': agrup_container,
            'combo': None,
            'campo_actual': None
        }

        campo_combo = customtkinter.CTkComboBox(
            agrup_container,
            values=nombres_campos if nombres_campos else ["Sin campos"],
            width=220,
            command=lambda choice: self._on_agrupacion_campo_change(agrup_obj, choice)
        )
        campo_combo.grid(row=0, column=2, sticky="w", padx=(0, 20))
        agrup_obj['combo'] = campo_combo

        # Botón eliminar
        del_btn = customtkinter.CTkButton(
            agrup_container,
            text="🗑️",
            width=40,
            fg_color="darkred",
            hover_color="red",
            command=lambda: self._remove_agrupacion(agrup_container)
        )
        del_btn.grid(row=0, column=3, padx=(0, 10))

        # Añadir a la lista
        self.agrupaciones.append(agrup_obj)

        # Auto-seleccionar primer campo
        if nombres_campos:
            campo_combo.set(nombres_campos[0])
            self._on_agrupacion_campo_change(agrup_obj, nombres_campos[0])

    def _on_agrupacion_campo_change(self, agrup_obj, campo_nombre):
        """Maneja el cambio de campo en una agrupación"""
        if not self.definicion_actual:
            return

        campos_informe = self.definicion_actual.get('campos', {})

        # Buscar el campo_key que corresponde al nombre seleccionado
        for campo_key, campo_def in campos_informe.items():
            if campo_def['nombre'] == campo_nombre:
                agrup_obj['campo_actual'] = campo_key
                break

    def _remove_agrupacion(self, container):
        """Elimina un nivel de agrupación"""
        container.destroy()
        # Actualizar lista
        self.agrupaciones = [a for a in self.agrupaciones if a['container'].winfo_exists()]

    def _add_agregacion(self):
        """Añade una nueva agregación (función de cálculo)"""
        if not self.definicion_actual:
            from CTkMessagebox import CTkMessagebox
            CTkMessagebox(
                title="Aviso",
                message="Selecciona primero un informe para poder configurar agregaciones",
                icon="warning"
            )
            return

        row = len(self.agregaciones)

        agreg_container = customtkinter.CTkFrame(self.agregaciones_frame, fg_color="gray25")
        agreg_container.grid(row=row, column=0, sticky="ew", pady=5, padx=5)
        agreg_container.grid_columnconfigure(1, weight=1)
        agreg_container.grid_columnconfigure(3, weight=1)

        # Label
        label = customtkinter.CTkLabel(
            agreg_container,
            text=f"Agregación {row + 1}:",
            font=customtkinter.CTkFont(size=11)
        )
        label.grid(row=0, column=0, padx=(10, 10), sticky="w")

        # Función
        funcion_label = customtkinter.CTkLabel(agreg_container, text="Función:")
        funcion_label.grid(row=0, column=1, sticky="w", padx=(0, 5))

        agregaciones_config = self.definicion_actual.get('agregaciones', {})
        funciones = list(agregaciones_config.keys())

        # Crear objeto de agregación
        agreg_obj = {
            'container': agreg_container,
            'funcion_combo': None,
            'campo_combo': None,
            'funcion_actual': None,
            'campo_actual': None
        }

        funcion_combo = customtkinter.CTkComboBox(
            agreg_container,
            values=funciones if funciones else ["COUNT"],
            width=150,
            command=lambda choice: self._on_agregacion_funcion_change(agreg_obj, choice)
        )
        funcion_combo.grid(row=0, column=2, sticky="w", padx=(0, 20))
        agreg_obj['funcion_combo'] = funcion_combo

        # Campo
        campo_label = customtkinter.CTkLabel(agreg_container, text="Campo:")
        campo_label.grid(row=0, column=3, sticky="w", padx=(0, 5))

        # Inicialmente vacío, se llenará cuando seleccione función
        campo_combo = customtkinter.CTkComboBox(
            agreg_container,
            values=["(Selecciona función)"],
            width=180,
            state="disabled"
        )
        campo_combo.grid(row=0, column=4, sticky="w", padx=(0, 10))
        agreg_obj['campo_combo'] = campo_combo

        # Botón eliminar
        del_btn = customtkinter.CTkButton(
            agreg_container,
            text="🗑️",
            width=40,
            fg_color="darkred",
            hover_color="red",
            command=lambda: self._remove_agregacion(agreg_container)
        )
        del_btn.grid(row=0, column=5, padx=(0, 10))

        # Añadir a la lista
        self.agregaciones.append(agreg_obj)

        # Auto-seleccionar primera función
        if funciones:
            funcion_combo.set(funciones[0])
            self._on_agregacion_funcion_change(agreg_obj, funciones[0])

    def _on_agregacion_funcion_change(self, agreg_obj, funcion):
        """Maneja el cambio de función en una agregación"""
        if not self.definicion_actual:
            return

        agreg_obj['funcion_actual'] = funcion

        # Obtener campos aplicables según la función
        agregaciones_config = self.definicion_actual.get('agregaciones', {})
        funcion_config = agregaciones_config.get(funcion, {})
        aplicable_a = funcion_config.get('aplicable_a', [])

        campos_informe = self.definicion_actual.get('campos', {})

        # Si aplica a todo (*), como COUNT, ofrecer "(Todos los registros)" + campos opcionales
        if "*" in aplicable_a:
            # COUNT puede aplicarse a todos los registros o a un campo específico
            nombres_campos = ["(Todos los registros)"]

            # Agregar también todos los campos disponibles como opciones
            for campo_key, campo_def in campos_informe.items():
                nombres_campos.append(campo_def['nombre'])

            # Habilitar el desplegable para que el usuario pueda elegir
            agreg_obj['campo_combo'].configure(values=nombres_campos, state="normal")
            agreg_obj['campo_combo'].set(nombres_campos[0])
            agreg_obj['campo_actual'] = None  # None significa COUNT(*)
        else:
            # Filtrar campos según el tipo
            nombres_campos = []
            for campo_key, campo_def in campos_informe.items():
                tipo_campo = campo_def.get('tipo')
                if tipo_campo in aplicable_a:
                    nombres_campos.append(campo_def['nombre'])

            if nombres_campos:
                agreg_obj['campo_combo'].configure(values=nombres_campos, state="normal")
                agreg_obj['campo_combo'].set(nombres_campos[0])
                # Guardar campo_key
                for campo_key, campo_def in campos_informe.items():
                    if campo_def['nombre'] == nombres_campos[0]:
                        agreg_obj['campo_actual'] = campo_key
                        break
            else:
                agreg_obj['campo_combo'].configure(values=["(No aplicable)"], state="disabled")
                agreg_obj['campo_combo'].set("(No aplicable)")
                agreg_obj['campo_actual'] = None

        # Añadir comando para detectar cambios de campo
        agreg_obj['campo_combo'].configure(
            command=lambda choice: self._on_agregacion_campo_change(agreg_obj, choice)
        )

    def _on_agregacion_campo_change(self, agreg_obj, campo_nombre):
        """Maneja el cambio de campo en una agregación"""
        if not self.definicion_actual:
            return

        # Si selecciona "(Todos los registros)", significa COUNT(*)
        if campo_nombre == "(Todos los registros)":
            agreg_obj['campo_actual'] = None
            return

        campos_informe = self.definicion_actual.get('campos', {})

        # Buscar el campo_key que corresponde al nombre seleccionado
        for campo_key, campo_def in campos_informe.items():
            if campo_def['nombre'] == campo_nombre:
                agreg_obj['campo_actual'] = campo_key
                break

    def _remove_agregacion(self, container):
        """Elimina una agregación"""
        container.destroy()
        # Actualizar lista
        self.agregaciones = [a for a in self.agregaciones if a['container'].winfo_exists()]

    def _on_modo_changed(self, modo):
        """Maneja el cambio de modo de visualización (Detalle/Resumen)"""
        self.modo_visualizacion = modo.lower()

    def _add_filtro(self):
        """Añade un nuevo selector de filtro dinámico"""
        if not self.definicion_actual:
            from CTkMessagebox import CTkMessagebox
            CTkMessagebox(
                title="Aviso",
                message="Selecciona primero un informe con filtros disponibles",
                icon="warning"
            )
            return

        row = len(self.filtros)

        filtro_container = customtkinter.CTkFrame(self.filtros_frame)
        filtro_container.grid(row=row, column=0, sticky="ew", pady=5)
        filtro_container.grid_columnconfigure(2, weight=1)

        # Lógica (Y/O) - solo si no es el primero
        if row > 0:
            logica_combo = customtkinter.CTkComboBox(
                filtro_container,
                values=LOGICA_FILTROS,
                width=60
            )
            logica_combo.grid(row=0, column=0, padx=(0, 10), sticky="w")
        else:
            logica_combo = None

        # Label
        label = customtkinter.CTkLabel(
            filtro_container,
            text=f"Filtro {row + 1}:",
            font=customtkinter.CTkFont(size=11)
        )
        label.grid(row=0, column=1, padx=(0, 10), sticky="w")

        # Campo - Poblar con filtros disponibles del informe
        campo_label = customtkinter.CTkLabel(filtro_container, text="Campo:")
        campo_label.grid(row=0, column=2, sticky="w", padx=(0, 5))

        filtros_disponibles = self.definicion_actual.get('filtros', {})
        nombres_campos = [self.definicion_actual['campos'][f['campo']]['nombre']
                         for f_key, f in filtros_disponibles.items()]

        campo_combo = customtkinter.CTkComboBox(
            filtro_container,
            values=nombres_campos if nombres_campos else ["Sin filtros"],
            width=150,
            command=lambda choice: self._on_filtro_campo_change(filtro_obj, choice)
        )
        campo_combo.grid(row=0, column=3, sticky="w", padx=(0, 15))

        # Operador
        operador_label = customtkinter.CTkLabel(filtro_container, text="Operador:")
        operador_label.grid(row=0, column=4, sticky="w", padx=(0, 5))

        operador_combo = customtkinter.CTkComboBox(
            filtro_container,
            values=["Seleccionar..."],
            width=130,
            command=lambda choice: self._on_filtro_operador_change(filtro_obj, choice)
        )
        operador_combo.grid(row=0, column=5, sticky="w", padx=(0, 15))

        # Valor - Widget inicial (se cambiará dinámicamente)
        valor_label = customtkinter.CTkLabel(filtro_container, text="Valor:")
        valor_label.grid(row=0, column=6, sticky="w", padx=(0, 5))

        valor_widget = customtkinter.CTkEntry(filtro_container, width=150)
        valor_widget.grid(row=0, column=7, sticky="w", padx=(0, 10))

        # Botón eliminar
        del_btn = customtkinter.CTkButton(
            filtro_container,
            text="🗑️",
            width=40,
            fg_color="darkred",
            hover_color="red",
            command=lambda: self._remove_filtro(filtro_container)
        )
        del_btn.grid(row=0, column=8, padx=(0, 5))

        # Objeto del filtro
        filtro_obj = {
            'container': filtro_container,
            'logica_combo': logica_combo,
            'campo_combo': campo_combo,
            'operador_combo': operador_combo,
            'valor_widget': valor_widget,
            'valor_label': valor_label,
            'campo_actual': None,
            'tipo_actual': None
        }

        self.filtros.append(filtro_obj)

        # Si hay campos, seleccionar el primero automáticamente
        if nombres_campos:
            campo_combo.set(nombres_campos[0])
            self._on_filtro_campo_change(filtro_obj, nombres_campos[0])

    def _on_filtro_campo_change(self, filtro_obj, campo_nombre):
        """Maneja el cambio de campo en un filtro"""
        if not self.definicion_actual:
            return

        # Buscar la definición del filtro por nombre del campo
        filtros_disponibles = self.definicion_actual.get('filtros', {})
        campos_def = self.definicion_actual.get('campos', {})

        # Encontrar el filtro correspondiente
        filtro_config = None
        campo_key = None
        for fkey, fconfig in filtros_disponibles.items():
            campo_def = campos_def.get(fconfig['campo'])
            if campo_def and campo_def['nombre'] == campo_nombre:
                filtro_config = fconfig
                campo_key = fconfig['campo']
                break

        if not filtro_config:
            return

        # Actualizar operadores disponibles
        operadores = filtro_config.get('operadores', [])
        filtro_obj['operador_combo'].configure(values=operadores)
        if operadores:
            filtro_obj['operador_combo'].set(operadores[0])

        # Almacenar tipo y campo actual
        filtro_obj['campo_actual'] = campo_key
        filtro_obj['tipo_actual'] = filtro_config.get('tipo')

        # Actualizar widget de valor según tipo
        self._update_valor_widget(filtro_obj, filtro_config)

    def _on_filtro_operador_change(self, filtro_obj, operador):
        """Maneja el cambio de operador (ajusta el widget de valor para casos especiales como 'Entre')"""
        # Si el operador es "Entre", mostrar dos campos de entrada
        if operador == "Entre":
            self._create_range_widget(filtro_obj)
        else:
            # Para otros operadores, usar el widget normal según el tipo de filtro
            if filtro_obj.get('tipo_actual') and self.definicion_actual:
                filtros_disponibles = self.definicion_actual.get('filtros', {})
                campo_key = filtro_obj.get('campo_actual')

                for fkey, fconfig in filtros_disponibles.items():
                    if fconfig['campo'] == campo_key:
                        self._update_valor_widget(filtro_obj, fconfig)
                        break

    def _create_range_widget(self, filtro_obj):
        """Crea dos campos de entrada para el operador 'Entre'"""
        # Destruir widget actual
        if filtro_obj.get('valor_widget'):
            filtro_obj['valor_widget'].destroy()
        if filtro_obj.get('valor_widget2'):
            filtro_obj['valor_widget2'].destroy()

        # Frame para contener ambos campos
        range_frame = customtkinter.CTkFrame(filtro_obj['container'], fg_color="transparent")
        range_frame.grid(row=0, column=7, sticky="w", padx=(0, 10))

        # Detectar si es un campo de fecha
        tipo_actual = filtro_obj.get('tipo_actual')

        if tipo_actual == 'fecha':
            # Para fechas, usar DateEntry
            widget1 = DateEntry(
                range_frame,
                width=11,
                background='darkblue',
                foreground='white',
                borderwidth=2,
                date_pattern='yyyy-mm-dd',
                locale='es_ES'
            )
            widget1.grid(row=0, column=0, padx=(0, 5))

            # Label "y"
            label_y = customtkinter.CTkLabel(range_frame, text="y", width=15)
            label_y.grid(row=0, column=1, padx=(0, 5))

            widget2 = DateEntry(
                range_frame,
                width=11,
                background='darkblue',
                foreground='white',
                borderwidth=2,
                date_pattern='yyyy-mm-dd',
                locale='es_ES'
            )
            widget2.grid(row=0, column=2)
        else:
            # Para numéricos, usar Entry normal
            widget1 = customtkinter.CTkEntry(
                range_frame,
                width=70,
                placeholder_text="Min..."
            )
            widget1.grid(row=0, column=0, padx=(0, 5))

            # Label "y"
            label_y = customtkinter.CTkLabel(range_frame, text="y", width=15)
            label_y.grid(row=0, column=1, padx=(0, 5))

            widget2 = customtkinter.CTkEntry(
                range_frame,
                width=70,
                placeholder_text="Max..."
            )
            widget2.grid(row=0, column=2)

        # Guardar ambos widgets
        filtro_obj['valor_widget'] = widget1
        filtro_obj['valor_widget2'] = widget2
        filtro_obj['is_range'] = True

    def _update_valor_widget(self, filtro_obj, filtro_config):
        """Actualiza el widget de valor según el tipo de filtro"""
        tipo = filtro_config.get('tipo')

        # Destruir widget actual
        if filtro_obj['valor_widget']:
            filtro_obj['valor_widget'].destroy()

        # Crear nuevo widget según tipo
        if tipo == 'select':
            # ComboBox con valores predefinidos
            valores = filtro_config.get('valores', [])
            widget = customtkinter.CTkComboBox(
                filtro_obj['container'],
                values=valores,
                width=150
            )

        elif tipo == 'select_bd':
            # ComboBox con valores de BD (obtener dinámicamente)
            tabla_dimension = filtro_config.get('tabla')
            valores = self._get_valores_dimension(tabla_dimension)
            widget = customtkinter.CTkComboBox(
                filtro_obj['container'],
                values=valores if valores else ["Sin datos"],
                width=150
            )

        elif tipo == 'numerico':
            # Entry numérico
            widget = customtkinter.CTkEntry(
                filtro_obj['container'],
                width=150,
                placeholder_text="Número..."
            )

        elif tipo == 'fecha':
            # DateEntry - selector de calendario
            widget = DateEntry(
                filtro_obj['container'],
                width=18,
                background='darkblue',
                foreground='white',
                borderwidth=2,
                date_pattern='yyyy-mm-dd',
                locale='es_ES'
            )

        else:
            # Default: Entry de texto
            widget = customtkinter.CTkEntry(filtro_obj['container'], width=150)

        widget.grid(row=0, column=7, sticky="w", padx=(0, 10))
        filtro_obj['valor_widget'] = widget

    def _get_valores_dimension(self, tabla_dimension):
        """Obtiene valores de una dimensión desde la BD"""
        try:
            resultados = get_dimension_values(self.user, self.password, self.schema, tabla_dimension)
            # Retornar solo las descripciones
            return [desc for id, desc in resultados]
        except Exception as e:
            print(f"Error al obtener valores de dimensión {tabla_dimension}: {e}")
            return []

    def _remove_filtro(self, container):
        """Elimina un selector de filtro"""
        container.destroy()
        # Actualizar lista
        self.filtros = [f for f in self.filtros if f['container'].winfo_exists()]

    def _clear_all_filtros(self):
        """Elimina todos los filtros"""
        for filtro in self.filtros:
            if filtro['container'].winfo_exists():
                filtro['container'].destroy()
        self.filtros = []

    def _clear_all_ordenaciones(self):
        """Elimina todas las ordenaciones"""
        for clasif in self.ordenaciones:
            if clasif['container'].winfo_exists():
                clasif['container'].destroy()
        self.ordenaciones = []

    def _update_campos_disponibles(self):
        """Actualiza los campos disponibles según el informe seleccionado"""
        # Limpiar campos actuales
        for widget in self.campos_scrollable.winfo_children():
            widget.destroy()

        self.campos_seleccionados = {}

        # Si hay definición del informe, usar sus campos
        if self.definicion_actual:
            campos_informe = self.definicion_actual.get('campos', {})
            campos_default = self.definicion_actual.get('campos_default', [])

            # Agrupar campos por grupo
            campos_por_grupo = {}
            for campo_key, campo_def in campos_informe.items():
                grupo = campo_def.get('grupo', 'Otros')
                if grupo not in campos_por_grupo:
                    campos_por_grupo[grupo] = []
                campos_por_grupo[grupo].append((campo_key, campo_def['nombre']))

            # Crear checkboxes por grupos
            row = 0
            for grupo, campos in campos_por_grupo.items():
                # Título del grupo
                grupo_label = customtkinter.CTkLabel(
                    self.campos_scrollable,
                    text=f"{grupo}:",
                    font=customtkinter.CTkFont(size=12, weight="bold")
                )
                grupo_label.grid(row=row, column=0, columnspan=3, sticky="w", pady=(10, 5))
                row += 1

                # Checkboxes en 3 columnas
                col = 0
                for campo_key, campo_nombre in campos:
                    # Marcar por defecto si está en campos_default
                    por_defecto = campo_key in campos_default
                    var = customtkinter.BooleanVar(value=por_defecto)
                    check = customtkinter.CTkCheckBox(
                        self.campos_scrollable,
                        text=campo_nombre,
                        variable=var
                    )
                    check.grid(row=row, column=col, sticky="w", padx=5, pady=2)

                    self.campos_seleccionados[campo_key] = var

                    col += 1
                    if col >= 3:
                        col = 0
                        row += 1

                if col > 0:
                    row += 1

        else:
            # Si no hay definición, usar campos genéricos de la categoría (fallback)
            campos_dict = None
            if "Partes" in str(self.categoria_seleccionada):
                campos_dict = CAMPOS_PARTES
            elif "Recursos" in str(self.categoria_seleccionada):
                campos_dict = CAMPOS_RECURSOS
            elif "Presupuestos" in str(self.categoria_seleccionada):
                campos_dict = CAMPOS_PRESUPUESTOS
            elif "Certificaciones" in str(self.categoria_seleccionada):
                campos_dict = CAMPOS_CERTIFICACIONES
            elif "Planificación" in str(self.categoria_seleccionada):
                campos_dict = CAMPOS_PLANIFICACION

            if campos_dict:
                row = 0
                for grupo, campos in campos_dict.items():
                    # Título del grupo
                    grupo_label = customtkinter.CTkLabel(
                        self.campos_scrollable,
                        text=f"{grupo}:",
                        font=customtkinter.CTkFont(size=12, weight="bold")
                    )
                    grupo_label.grid(row=row, column=0, columnspan=3, sticky="w", pady=(10, 5))
                    row += 1

                    # Checkboxes en 3 columnas
                    col = 0
                    for campo in campos:
                        var = customtkinter.BooleanVar(value=True)
                        check = customtkinter.CTkCheckBox(
                            self.campos_scrollable,
                            text=campo,
                            variable=var
                        )
                        check.grid(row=row, column=col, sticky="w", padx=5, pady=2)

                        self.campos_seleccionados[campo] = var

                        col += 1
                        if col >= 3:
                            col = 0
                            row += 1

                    if col > 0:
                        row += 1

    def _select_all_campos(self):
        """Selecciona todos los campos"""
        for var in self.campos_seleccionados.values():
            var.set(True)

    def _deselect_all_campos(self):
        """Deselecciona todos los campos"""
        for var in self.campos_seleccionados.values():
            var.set(False)

    def _open_config_dialog(self):
        """Abre el diálogo de configuración de cabecera"""
        from CTkMessagebox import CTkMessagebox

        CTkMessagebox(
            title="Configuración",
            message="Funcionalidad de configuración de cabecera en desarrollo.\n\n"
                    "Próximamente podrás configurar:\n"
                    "• Datos de la empresa\n"
                    "• Logo corporativo\n"
                    "• Datos del proyecto\n"
                    "• Pie de página personalizado",
            icon="info"
        )

    def _preview_report(self):
        """Previsualiza el informe ejecutando el query y mostrando resultados"""
        from CTkMessagebox import CTkMessagebox

        # Validaciones
        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Por favor, seleccione un informe del menú izquierdo.",
                icon="warning"
            )
            return

        if not self.definicion_actual:
            CTkMessagebox(
                title="Aviso",
                message=f"El informe '{self.informe_seleccionado}' aún no está implementado.",
                icon="warning"
            )
            return

        # Recopilar filtros aplicados
        filtros_aplicados = []
        for filtro_obj in self.filtros:
            campo_actual = filtro_obj.get('campo_actual')
            if not campo_actual:
                continue

            operador = filtro_obj['operador_combo'].get()
            valor_widget = filtro_obj['valor_widget']

            # Obtener valor según tipo de widget
            # Caso especial: operador "Entre" requiere dos valores
            if filtro_obj.get('is_range') and operador == "Entre":
                widget1 = filtro_obj.get('valor_widget')
                widget2 = filtro_obj.get('valor_widget2')

                if widget1 and widget2:
                    # Obtener valor según tipo de widget
                    if isinstance(widget1, (customtkinter.CTkEntry, DateEntry)):
                        valor1 = widget1.get()
                    else:
                        valor1 = ""

                    if isinstance(widget2, (customtkinter.CTkEntry, DateEntry)):
                        valor2 = widget2.get()
                    else:
                        valor2 = ""

                    if not valor1 or not valor2:
                        continue

                    # Para "Entre", pasar tupla (min, max)
                    valor = (valor1, valor2)
                else:
                    continue
            else:
                # Caso normal: un solo valor
                if isinstance(valor_widget, customtkinter.CTkComboBox):
                    valor = valor_widget.get()
                elif isinstance(valor_widget, (customtkinter.CTkEntry, DateEntry)):
                    valor = valor_widget.get()
                else:
                    valor = ""

            if not valor or valor == "Seleccionar..." or not operador or operador == "Seleccionar...":
                continue

            # Obtener lógica (Y/O) del combo - por defecto 'Y'
            logica = 'Y'
            if filtro_obj.get('logica_combo'):
                logica = filtro_obj['logica_combo'].get()

            filtros_aplicados.append({
                'campo': campo_actual,
                'operador': operador,
                'valor': valor,
                'logica': logica
            })

        # Recopilar ordenaciones aplicadas
        ordenaciones_aplicadas = []
        for clasif_obj in self.ordenaciones:
            campo_actual = clasif_obj.get('campo_actual')
            if not campo_actual:
                continue

            orden = clasif_obj['orden_combo'].get()
            if not orden or orden == "Seleccionar...":
                orden = "Ascendente"  # Valor por defecto

            ordenaciones_aplicadas.append({
                'campo': campo_actual,
                'orden': orden
            })

        # Recopilar campos seleccionados
        campos_seleccionados = [campo_key for campo_key, var in self.campos_seleccionados.items() if var.get()]

        if not campos_seleccionados:
            CTkMessagebox(
                title="Aviso",
                message="Seleccione al menos un campo para mostrar en el informe.",
                icon="warning"
            )
            return

        # Recopilar agrupaciones aplicadas
        agrupaciones_aplicadas = []
        for agrup_obj in self.agrupaciones:
            campo_actual = agrup_obj.get('campo_actual')
            if campo_actual:
                agrupaciones_aplicadas.append(campo_actual)

        # Recopilar agregaciones aplicadas
        agregaciones_aplicadas = []
        for agreg_obj in self.agregaciones:
            funcion = agreg_obj.get('funcion_actual')
            campo = agreg_obj.get('campo_actual')  # Puede ser None para COUNT
            if funcion:
                agregaciones_aplicadas.append({
                    'funcion': funcion,
                    'campo': campo
                })

        # Ejecutar informe
        print(f"\n{'='*70}")
        print(f"EJECUTANDO INFORME: {self.informe_seleccionado}")
        print(f"Filtros aplicados: {len(filtros_aplicados)}")
        print(f"Ordenaciones aplicadas: {len(ordenaciones_aplicadas)}")
        print(f"Campos seleccionados: {len(campos_seleccionados)}")
        print(f"Agrupaciones aplicadas: {len(agrupaciones_aplicadas)}")
        print(f"Agregaciones aplicadas: {len(agregaciones_aplicadas)}")
        print(f"Modo visualización: {self.modo_visualizacion}")
        print(f"{'='*70}\n")

        try:
            # Si hay agrupaciones o agregaciones, usar la versión extendida
            if agrupaciones_aplicadas or agregaciones_aplicadas:
                columnas, datos, resultado_agrupacion = ejecutar_informe_con_agrupacion(
                    self.user,
                    self.password,
                    self.schema,
                    self.informe_seleccionado,
                    filtros=filtros_aplicados,
                    ordenaciones=ordenaciones_aplicadas,
                    campos_seleccionados=campos_seleccionados,
                    agrupaciones=agrupaciones_aplicadas,
                    agregaciones=agregaciones_aplicadas,
                    modo=self.modo_visualizacion
                )
                # Convertir resultado_agrupacion a totales para compatibilidad
                totales = resultado_agrupacion.get('totales_generales', {})
            else:
                columnas, datos, totales = ejecutar_informe(
                    self.user,
                    self.password,
                    self.schema,
                    self.informe_seleccionado,
                    filtros=filtros_aplicados,
                    ordenaciones=ordenaciones_aplicadas,
                    campos_seleccionados=campos_seleccionados
                )
                resultado_agrupacion = None

            # Mostrar resultados
            if datos:
                self._show_results_window(columnas, datos, totales, resultado_agrupacion)
            else:
                # Mensaje más claro dependiendo de si hay filtros o no
                if filtros_aplicados:
                    mensaje = "No se encontraron datos con los filtros aplicados.\n\nIntente modificar o eliminar algunos filtros."
                else:
                    mensaje = "No se encontraron datos en la base de datos para este informe.\n\nVerifique que la tabla contenga registros."

                CTkMessagebox(
                    title="Sin Resultados",
                    message=mensaje,
                    icon="info"
                )

        except Exception as e:
            import traceback
            error_msg = str(e)
            print(f"Error al generar informe:\n{traceback.format_exc()}")

            CTkMessagebox(
                title="Error",
                message=f"Error al generar el informe:\n\n{error_msg}",
                icon="cancel"
            )

    def _crear_fila_totales(self, columnas, totales_dict, texto_primera_col="═══ TOTAL ═══"):
        """Crea una fila de totales formateada

        Args:
            columnas: Lista de nombres de columnas
            totales_dict: Dict con los totales {función(campo): valor}
            texto_primera_col: Texto para la primera columna

        Returns:
            Lista de valores formateados para la fila de totales
        """
        fila = []
        for i, col in enumerate(columnas):
            # Buscar si hay un total para esta columna
            total_encontrado = False
            for key, valor in totales_dict.items():
                # key tiene formato "FUNCION(campo)" o "FUNCION(*)"
                if f"({col})" in key or (col in totales_dict):
                    fila.append(f"{valor:,.2f}" if isinstance(valor, (int, float)) else str(valor))
                    total_encontrado = True
                    break

            if not total_encontrado:
                if i == 0:
                    fila.append(texto_primera_col)
                else:
                    fila.append("")

        return fila

    def _render_grupos_recursivo(self, tree, grupos, columnas, parent="", nivel=0, modo='detalle'):
        """Renderiza grupos de forma recursiva en el TreeView

        Args:
            tree: Widget TreeView
            grupos: Lista de grupos a renderizar
            columnas: Lista de nombres de columnas
            parent: ID del nodo padre ("" para raíz)
            nivel: Nivel de profundidad del grupo (0 = primer nivel)
            modo: 'detalle' o 'resumen'
        """
        if not grupos:
            return

        for grupo in grupos:
            clave = grupo['clave']
            campo = grupo['campo']
            datos = grupo['datos']
            subtotales = grupo.get('subtotales', {})
            subgrupos = grupo.get('subgrupos')

            # ENCABEZADO DEL GRUPO
            # Crear fila de encabezado con el nombre del grupo (sin subtotales)
            indent = "    " * nivel  # Indentación visual
            titulo_grupo = f"{indent}📁 {campo.upper()}: {clave}"

            # Crear fila de encabezado
            fila_header = [titulo_grupo] + [""] * (len(columnas) - 1)
            header_id = tree.insert(parent, "end", values=fila_header, tags=(f'grupo_header_nivel{nivel}',))

            # Configurar estilo según nivel
            if nivel == 0:
                tree.tag_configure(f'grupo_header_nivel{nivel}', background='#4A6FA5', foreground='white', font=('TkDefaultFont', 10, 'bold'))
            elif nivel == 1:
                tree.tag_configure(f'grupo_header_nivel{nivel}', background='#6B8FB8', foreground='white', font=('TkDefaultFont', 9, 'bold'))
            else:
                tree.tag_configure(f'grupo_header_nivel{nivel}', background='#8AADC7', foreground='white', font=('TkDefaultFont', 9))

            # SUBGRUPOS (si existen)
            if subgrupos:
                self._render_grupos_recursivo(tree, subgrupos, columnas, header_id, nivel + 1, modo)
            elif modo == 'detalle':
                # DATOS DEL GRUPO (solo si estamos en modo detalle y no hay subgrupos)
                for fila_datos in datos:
                    tree.insert(header_id, "end", values=fila_datos, tags=(f'datos_nivel{nivel}',))

                # Estilo para datos
                if nivel == 0:
                    tree.tag_configure(f'datos_nivel{nivel}', background='gray20')
                else:
                    tree.tag_configure(f'datos_nivel{nivel}', background='gray15')

            # FILA DE SUBTOTALES (si hay agregaciones configuradas)
            if subtotales:
                # Crear fila de subtotales con valores en las columnas correspondientes
                indent_subtotal = "    " * (nivel + 1)
                fila_subtotal = []

                for i, col_name in enumerate(columnas):
                    if i == 0:
                        # Primera columna: etiqueta de subtotal
                        fila_subtotal.append(f"{indent_subtotal}▸ Subtotal")
                    else:
                        # Buscar si hay un subtotal para esta columna
                        valor_subtotal = ""
                        for key, valor in subtotales.items():
                            # Las claves son del tipo "SUM(presupuesto)", "COUNT(*)", etc.
                            # Necesitamos extraer el nombre del campo
                            if "(" in key and ")" in key:
                                # Extraer campo de la función: "SUM(presupuesto)" -> "presupuesto"
                                campo_agg = key.split("(")[1].split(")")[0]
                                if campo_agg == col_name or campo_agg == "*":
                                    # Formatear el valor
                                    if isinstance(valor, (int, float)):
                                        if isinstance(valor, float):
                                            valor_subtotal = f"{valor:,.2f}"
                                        else:
                                            valor_subtotal = f"{valor:,}"
                                    else:
                                        valor_subtotal = str(valor)
                                    break

                        fila_subtotal.append(valor_subtotal)

                # Insertar fila de subtotales
                subtotal_id = tree.insert(header_id, "end", values=fila_subtotal, tags=(f'subtotal_nivel{nivel}',))

                # Configurar estilo para subtotales
                if nivel == 0:
                    tree.tag_configure(f'subtotal_nivel{nivel}', background='#2C5F8D', foreground='white', font=('TkDefaultFont', 9, 'bold'))
                elif nivel == 1:
                    tree.tag_configure(f'subtotal_nivel{nivel}', background='#4A7BA7', foreground='white', font=('TkDefaultFont', 9, 'bold'))
                else:
                    tree.tag_configure(f'subtotal_nivel{nivel}', background='#6B95BC', foreground='white', font=('TkDefaultFont', 9))

    def _show_results_window(self, columnas, datos, totales=None, resultado_agrupacion=None):
        """Muestra una ventana con los resultados del informe

        Args:
            columnas: Lista de nombres de columnas
            datos: Lista de tuplas con los datos
            totales: Dict con totales por columna
            resultado_agrupacion: Dict con estructura de grupos y agregaciones (opcional)
        """
        # Crear ventana toplevel
        results_window = customtkinter.CTkToplevel(self)
        results_window.title(f"Vista Previa: {self.informe_seleccionado}")
        results_window.geometry("1200x600")

        # Frame principal
        main_frame = customtkinter.CTkFrame(results_window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(1, weight=1)

        # Título
        title_label = customtkinter.CTkLabel(
            main_frame,
            text=f"📊 {self.informe_seleccionado}",
            font=customtkinter.CTkFont(size=16, weight="bold")
        )
        title_label.grid(row=0, column=0, pady=(5, 10), sticky="w", padx=10)

        # Info
        info_label = customtkinter.CTkLabel(
            main_frame,
            text=f"{len(datos)} registros encontrados",
            font=customtkinter.CTkFont(size=12),
            text_color="gray"
        )
        info_label.grid(row=0, column=1, pady=(5, 10), sticky="e", padx=10)

        # Frame para TreeView
        tree_frame = customtkinter.CTkFrame(main_frame)
        tree_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=(0, 10))

        # Scrollbars
        vsb = customtkinter.CTkScrollbar(tree_frame, orientation="vertical")
        vsb.pack(side="right", fill="y")

        hsb = customtkinter.CTkScrollbar(tree_frame, orientation="horizontal")
        hsb.pack(side="bottom", fill="x")

        # TreeView
        tree = ttk.Treeview(
            tree_frame,
            columns=columnas,
            show="headings",
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set
        )
        tree.pack(side="left", fill="both", expand=True)

        vsb.configure(command=tree.yview)
        hsb.configure(command=tree.xview)

        # Configurar columnas
        for col in columnas:
            tree.heading(col, text=col)
            tree.column(col, width=150, anchor="w")

        # Insertar datos con soporte para agrupaciones
        if resultado_agrupacion and resultado_agrupacion.get('grupos'):
            # MODO AGRUPADO: Renderizar grupos jerárquicos
            modo = resultado_agrupacion.get('modo', 'detalle')
            self._render_grupos_recursivo(tree, resultado_agrupacion['grupos'], columnas, "", nivel=0, modo=modo)

            # Añadir fila de totales generales al final
            if resultado_agrupacion.get('totales_generales') and self.totales_var.get():
                tree.insert("", "end", values=[""] * len(columnas), tags=('separador',))
                fila_totales = self._crear_fila_totales(
                    columnas,
                    resultado_agrupacion['totales_generales'],
                    "═══ TOTAL GENERAL ═══"
                )
                tree.insert("", "end", values=fila_totales, tags=('total_general',))
                tree.tag_configure('separador', background='gray40')
                tree.tag_configure('total_general', background='#2E5C8A', foreground='white', font=('TkDefaultFont', 10, 'bold'))
        else:
            # MODO NORMAL: Insertar datos sin agrupación
            for fila in datos:
                tree.insert("", "end", values=fila)

        # Insertar fila de totales si hay totales y la opción está activada (modo normal)
        if not resultado_agrupacion and totales and self.totales_var.get():
            # Crear fila de totales
            fila_totales = []
            for i, col in enumerate(columnas):
                if col in totales:
                    # Formatear el total según sea moneda o número
                    valor = totales[col]
                    fila_totales.append(f"{valor:,.2f} €" if valor else "0.00 €")
                elif i == 0:
                    # Primera columna: texto "TOTAL"
                    fila_totales.append("═══ TOTAL ═══")
                else:
                    # Otras columnas: vacío
                    fila_totales.append("")

            # Insertar fila de totales con tag especial
            total_item = tree.insert("", "end", values=fila_totales, tags=('total',))
            # Configurar estilo para fila de totales
            tree.tag_configure('total', background='#4472C4', foreground='white', font=('TkDefaultFont', 10, 'bold'))

        # Botón cerrar
        close_btn = customtkinter.CTkButton(
            main_frame,
            text="Cerrar",
            width=100,
            command=results_window.destroy
        )
        close_btn.grid(row=2, column=0, columnspan=2, pady=(0, 5))

        # Asegurar que la ventana se muestre al frente
        results_window.transient(self)  # Hacer ventana dependiente de la principal
        results_window.update_idletasks()

        # Forzar que aparezca al frente (especialmente en Windows)
        results_window.attributes('-topmost', True)
        results_window.lift()
        results_window.focus_force()

        # Después de 100ms, quitar el "siempre encima" para permitir navegación normal
        results_window.after(100, lambda: results_window.attributes('-topmost', False))

    def _export_word(self):
        """Exporta el informe a formato Word (.docx)"""
        from CTkMessagebox import CTkMessagebox
        from tkinter import filedialog
        import datetime

        # Validaciones
        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Por favor, seleccione un informe del menú izquierdo.",
                icon="warning"
            )
            return

        if not self.definicion_actual:
            CTkMessagebox(
                title="Aviso",
                message=f"El informe '{self.informe_seleccionado}' aún no está implementado.",
                icon="warning"
            )
            return

        # Recopilar filtros aplicados
        filtros_aplicados = []
        for filtro_obj in self.filtros:
            campo_actual = filtro_obj.get('campo_actual')
            if not campo_actual:
                continue

            operador = filtro_obj['operador_combo'].get()
            valor_widget = filtro_obj['valor_widget']

            if isinstance(valor_widget, customtkinter.CTkComboBox):
                valor = valor_widget.get()
            elif isinstance(valor_widget, customtkinter.CTkEntry):
                valor = valor_widget.get()
            else:
                valor = ""

            if not valor or valor == "Seleccionar..." or not operador or operador == "Seleccionar...":
                continue

            # Obtener lógica (Y/O) del combo - por defecto 'Y'
            logica = 'Y'
            if filtro_obj.get('logica_combo'):
                logica = filtro_obj['logica_combo'].get()

            filtros_aplicados.append({
                'campo': campo_actual,
                'operador': operador,
                'valor': valor,
                'logica': logica
            })

        # Recopilar ordenaciones aplicadas
        ordenaciones_aplicadas = []
        for clasif_obj in self.ordenaciones:
            campo_actual = clasif_obj.get('campo_actual')
            if not campo_actual:
                continue

            orden = clasif_obj['orden_combo'].get()
            if not orden or orden == "Seleccionar...":
                orden = "Ascendente"

            ordenaciones_aplicadas.append({
                'campo': campo_actual,
                'orden': orden
            })

        # Recopilar campos seleccionados
        campos_seleccionados = [campo_key for campo_key, var in self.campos_seleccionados.items() if var.get()]

        if not campos_seleccionados:
            CTkMessagebox(
                title="Aviso",
                message="Seleccione al menos un campo para incluir en el informe.",
                icon="warning"
            )
            return

        # Ejecutar informe para obtener datos
        try:
            columnas, datos, totales = ejecutar_informe(
                self.user,
                self.password,
                self.schema,
                self.informe_seleccionado,
                filtros=filtros_aplicados,
                ordenaciones=ordenaciones_aplicadas,
                campos_seleccionados=campos_seleccionados
            )

            if not datos:
                # Mensaje más claro dependiendo de si hay filtros o no
                if filtros_aplicados:
                    mensaje = "No se encontraron datos con los filtros aplicados.\n\nIntente modificar o eliminar algunos filtros."
                else:
                    mensaje = "No se encontraron datos en la base de datos para este informe.\n\nVerifique que la tabla contenga registros."

                CTkMessagebox(
                    title="Sin Resultados",
                    message=mensaje,
                    icon="info"
                )
                return

        except Exception as e:
            import traceback
            error_msg = str(e)
            print(f"Error al generar informe:\n{traceback.format_exc()}")

            CTkMessagebox(
                title="Error",
                message=f"Error al generar el informe:\n\n{error_msg}",
                icon="cancel"
            )
            return

        # Pedir ubicación de guardado
        fecha_actual = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{self.informe_seleccionado.replace(' ', '_')}_{fecha_actual}.docx"

        archivo = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialfile=nombre_archivo,
            title="Guardar informe como..."
        )

        if not archivo:
            return  # Usuario canceló

        # Crear archivo Word
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Título del informe
            titulo = doc.add_heading(f"📊 {self.informe_seleccionado}", level=0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titulo_run = titulo.runs[0]
            titulo_run.font.color.rgb = RGBColor(31, 78, 120)  # Azul oscuro

            # Información de fecha y filtros
            info_text = f"Generado: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
            if filtros_aplicados:
                info_text += f" | Filtros aplicados: {len(filtros_aplicados)}"
            if ordenaciones_aplicadas:
                info_text += f" | Ordenaciones: {len(ordenaciones_aplicadas)}"

            info_para = doc.add_paragraph(info_text)
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.runs[0]
            info_run.font.size = Pt(10)
            info_run.font.italic = True
            info_run.font.color.rgb = RGBColor(102, 102, 102)  # Gris

            # Espacio
            doc.add_paragraph()

            # Tabla con datos
            tabla = doc.add_table(rows=1, cols=len(columnas))
            tabla.style = 'Light Grid Accent 1'

            # Encabezados
            header_cells = tabla.rows[0].cells
            for idx, columna in enumerate(columnas):
                header_cells[idx].text = str(columna)
                # Formatear encabezado
                for paragraph in header_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Color de fondo se maneja con el estilo de tabla

            # Datos
            for fila in datos:
                row_cells = tabla.add_row().cells
                for idx, valor in enumerate(fila):
                    row_cells[idx].text = str(valor) if valor is not None else ""
                    # Formatear dato
                    for paragraph in row_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            # Espacio
            doc.add_paragraph()

            # Resumen
            resumen_para = doc.add_paragraph(f"Total de registros: {len(datos)}")
            resumen_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            resumen_run = resumen_para.runs[0]
            resumen_run.font.bold = True
            resumen_run.font.size = Pt(11)

            # Guardar archivo
            doc.save(archivo)

            CTkMessagebox(
                title="Exportación Exitosa",
                message=f"El informe se ha exportado correctamente a:\n\n{archivo}\n\n"
                        f"Registros exportados: {len(datos)}",
                icon="check"
            )

        except ImportError:
            CTkMessagebox(
                title="Error",
                message="La librería 'python-docx' no está instalada.\n\n"
                        "Por favor, instálala con:\n"
                        "pip install python-docx",
                icon="cancel"
            )
        except Exception as e:
            import traceback
            print(f"Error al exportar a Word:\n{traceback.format_exc()}")
            CTkMessagebox(
                title="Error",
                message=f"Error al exportar a Word:\n\n{str(e)}",
                icon="cancel"
            )

    def _export_excel(self):
        """Exporta el informe a formato Excel (.xlsx)"""
        from CTkMessagebox import CTkMessagebox
        from tkinter import filedialog
        import datetime

        # Validaciones
        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Por favor, seleccione un informe del menú izquierdo.",
                icon="warning"
            )
            return

        if not self.definicion_actual:
            CTkMessagebox(
                title="Aviso",
                message=f"El informe '{self.informe_seleccionado}' aún no está implementado.",
                icon="warning"
            )
            return

        # Recopilar filtros aplicados
        filtros_aplicados = []
        for filtro_obj in self.filtros:
            campo_actual = filtro_obj.get('campo_actual')
            if not campo_actual:
                continue

            operador = filtro_obj['operador_combo'].get()
            valor_widget = filtro_obj['valor_widget']

            # Obtener valor según tipo de widget
            # Caso especial: operador "Entre" requiere dos valores
            if filtro_obj.get('is_range') and operador == "Entre":
                widget1 = filtro_obj.get('valor_widget')
                widget2 = filtro_obj.get('valor_widget2')

                if widget1 and widget2:
                    # Obtener valor según tipo de widget
                    if isinstance(widget1, (customtkinter.CTkEntry, DateEntry)):
                        valor1 = widget1.get()
                    else:
                        valor1 = ""

                    if isinstance(widget2, (customtkinter.CTkEntry, DateEntry)):
                        valor2 = widget2.get()
                    else:
                        valor2 = ""

                    if not valor1 or not valor2:
                        continue

                    # Para "Entre", pasar tupla (min, max)
                    valor = (valor1, valor2)
                else:
                    continue
            else:
                # Caso normal: un solo valor
                if isinstance(valor_widget, customtkinter.CTkComboBox):
                    valor = valor_widget.get()
                elif isinstance(valor_widget, (customtkinter.CTkEntry, DateEntry)):
                    valor = valor_widget.get()
                else:
                    valor = ""

            if not valor or valor == "Seleccionar..." or not operador or operador == "Seleccionar...":
                continue

            # Obtener lógica (Y/O) del combo - por defecto 'Y'
            logica = 'Y'
            if filtro_obj.get('logica_combo'):
                logica = filtro_obj['logica_combo'].get()

            filtros_aplicados.append({
                'campo': campo_actual,
                'operador': operador,
                'valor': valor,
                'logica': logica
            })

        # Recopilar ordenaciones aplicadas
        ordenaciones_aplicadas = []
        for clasif_obj in self.ordenaciones:
            campo_actual = clasif_obj.get('campo_actual')
            if not campo_actual:
                continue

            orden = clasif_obj['orden_combo'].get()
            if not orden or orden == "Seleccionar...":
                orden = "Ascendente"

            ordenaciones_aplicadas.append({
                'campo': campo_actual,
                'orden': orden
            })

        # Recopilar campos seleccionados
        campos_seleccionados = [campo_key for campo_key, var in self.campos_seleccionados.items() if var.get()]

        if not campos_seleccionados:
            CTkMessagebox(
                title="Aviso",
                message="Seleccione al menos un campo para incluir en el informe.",
                icon="warning"
            )
            return

        # Ejecutar informe para obtener datos
        try:
            columnas, datos, totales = ejecutar_informe(
                self.user,
                self.password,
                self.schema,
                self.informe_seleccionado,
                filtros=filtros_aplicados,
                ordenaciones=ordenaciones_aplicadas,
                campos_seleccionados=campos_seleccionados
            )

            if not datos:
                # Mensaje más claro dependiendo de si hay filtros o no
                if filtros_aplicados:
                    mensaje = "No se encontraron datos con los filtros aplicados.\n\nIntente modificar o eliminar algunos filtros."
                else:
                    mensaje = "No se encontraron datos en la base de datos para este informe.\n\nVerifique que la tabla contenga registros."

                CTkMessagebox(
                    title="Sin Resultados",
                    message=mensaje,
                    icon="info"
                )
                return

        except Exception as e:
            import traceback
            error_msg = str(e)
            print(f"Error al generar informe:\n{traceback.format_exc()}")

            CTkMessagebox(
                title="Error",
                message=f"Error al generar el informe:\n\n{error_msg}",
                icon="cancel"
            )
            return

        # Pedir ubicación de guardado
        fecha_actual = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{self.informe_seleccionado.replace(' ', '_')}_{fecha_actual}.xlsx"

        archivo = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            initialfile=nombre_archivo,
            title="Guardar informe como..."
        )

        if not archivo:
            return  # Usuario canceló

        # Crear archivo Excel
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = self.informe_seleccionado[:31]  # Límite de 31 caracteres para nombre de hoja

            # Título del informe (fila 1)
            ws.merge_cells('A1:' + chr(64 + len(columnas)) + '1')
            titulo_cell = ws['A1']
            titulo_cell.value = f"📊 {self.informe_seleccionado}"
            titulo_cell.font = Font(size=16, bold=True, color="FFFFFF")
            titulo_cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            titulo_cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[1].height = 30

            # Información de fecha y filtros (fila 2)
            info_text = f"Generado: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
            if filtros_aplicados:
                info_text += f" | Filtros aplicados: {len(filtros_aplicados)}"
            ws.merge_cells('A2:' + chr(64 + len(columnas)) + '2')
            info_cell = ws['A2']
            info_cell.value = info_text
            info_cell.font = Font(size=10, italic=True, color="666666")
            info_cell.alignment = Alignment(horizontal="center")
            ws.row_dimensions[2].height = 20

            # Encabezados de columnas (fila 4)
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            for col_idx, columna in enumerate(columnas, start=1):
                cell = ws.cell(row=4, column=col_idx)
                cell.value = columna
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
                # Ajustar ancho de columna
                ws.column_dimensions[chr(64 + col_idx)].width = max(15, len(str(columna)) + 2)

            ws.row_dimensions[4].height = 25

            # Datos (desde fila 5)
            data_alignment = Alignment(horizontal="left", vertical="center")
            for row_idx, fila in enumerate(datos, start=5):
                for col_idx, valor in enumerate(fila, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.value = valor
                    cell.alignment = data_alignment
                    cell.border = border

                    # Alternar color de fila
                    if row_idx % 2 == 0:
                        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

            # Resumen al final
            ultima_fila = 5 + len(datos)
            ws.merge_cells(f'A{ultima_fila}:' + chr(64 + len(columnas)) + str(ultima_fila))
            resumen_cell = ws[f'A{ultima_fila}']
            resumen_cell.value = f"Total de registros: {len(datos)}"
            resumen_cell.font = Font(bold=True, size=11)
            resumen_cell.alignment = Alignment(horizontal="right")
            resumen_cell.fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")

            # Guardar archivo
            wb.save(archivo)

            CTkMessagebox(
                title="Exportación Exitosa",
                message=f"El informe se ha exportado correctamente a:\n\n{archivo}\n\n"
                        f"Registros exportados: {len(datos)}",
                icon="check"
            )

        except ImportError:
            CTkMessagebox(
                title="Error",
                message="La librería 'openpyxl' no está instalada.\n\n"
                        "Por favor, instálala con:\n"
                        "pip install openpyxl",
                icon="cancel"
            )
        except Exception as e:
            import traceback
            print(f"Error al exportar a Excel:\n{traceback.format_exc()}")
            CTkMessagebox(
                title="Error",
                message=f"Error al exportar a Excel:\n\n{str(e)}",
                icon="cancel"
            )

    def _export_pdf(self):
        """Exporta el informe a formato PDF"""
        from CTkMessagebox import CTkMessagebox
        from tkinter import filedialog
        import datetime

        # Validaciones
        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Por favor, seleccione un informe del menú izquierdo.",
                icon="warning"
            )
            return

        if not self.definicion_actual:
            CTkMessagebox(
                title="Aviso",
                message=f"El informe '{self.informe_seleccionado}' aún no está implementado.",
                icon="warning"
            )
            return

        # Recopilar filtros aplicados
        filtros_aplicados = []
        for filtro_obj in self.filtros:
            campo_actual = filtro_obj.get('campo_actual')
            if not campo_actual:
                continue

            operador = filtro_obj['operador_combo'].get()
            valor_widget = filtro_obj['valor_widget']

            if isinstance(valor_widget, customtkinter.CTkComboBox):
                valor = valor_widget.get()
            elif isinstance(valor_widget, customtkinter.CTkEntry):
                valor = valor_widget.get()
            else:
                valor = ""

            if not valor or valor == "Seleccionar..." or not operador or operador == "Seleccionar...":
                continue

            # Obtener lógica (Y/O) del combo - por defecto 'Y'
            logica = 'Y'
            if filtro_obj.get('logica_combo'):
                logica = filtro_obj['logica_combo'].get()

            filtros_aplicados.append({
                'campo': campo_actual,
                'operador': operador,
                'valor': valor,
                'logica': logica
            })

        # Recopilar ordenaciones aplicadas
        ordenaciones_aplicadas = []
        for clasif_obj in self.ordenaciones:
            campo_actual = clasif_obj.get('campo_actual')
            if not campo_actual:
                continue

            orden = clasif_obj['orden_combo'].get()
            if not orden or orden == "Seleccionar...":
                orden = "Ascendente"

            ordenaciones_aplicadas.append({
                'campo': campo_actual,
                'orden': orden
            })

        # Recopilar campos seleccionados
        campos_seleccionados = [campo_key for campo_key, var in self.campos_seleccionados.items() if var.get()]

        if not campos_seleccionados:
            CTkMessagebox(
                title="Aviso",
                message="Seleccione al menos un campo para incluir en el informe.",
                icon="warning"
            )
            return

        # Ejecutar informe para obtener datos
        try:
            columnas, datos, totales = ejecutar_informe(
                self.user,
                self.password,
                self.schema,
                self.informe_seleccionado,
                filtros=filtros_aplicados,
                ordenaciones=ordenaciones_aplicadas,
                campos_seleccionados=campos_seleccionados
            )

            if not datos:
                # Mensaje más claro dependiendo de si hay filtros o no
                if filtros_aplicados:
                    mensaje = "No se encontraron datos con los filtros aplicados.\n\nIntente modificar o eliminar algunos filtros."
                else:
                    mensaje = "No se encontraron datos en la base de datos para este informe.\n\nVerifique que la tabla contenga registros."

                CTkMessagebox(
                    title="Sin Resultados",
                    message=mensaje,
                    icon="info"
                )
                return

        except Exception as e:
            import traceback
            error_msg = str(e)
            print(f"Error al generar informe:\n{traceback.format_exc()}")

            CTkMessagebox(
                title="Error",
                message=f"Error al generar el informe:\n\n{error_msg}",
                icon="cancel"
            )
            return

        # Pedir ubicación de guardado
        fecha_actual = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{self.informe_seleccionado.replace(' ', '_')}_{fecha_actual}.pdf"

        archivo = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            initialfile=nombre_archivo,
            title="Guardar informe como..."
        )

        if not archivo:
            return  # Usuario canceló

        # Crear archivo PDF
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_RIGHT

            # Determinar orientación según número de columnas
            if len(columnas) > 6:
                pagesize = landscape(A4)
            else:
                pagesize = A4

            # Crear documento
            doc = SimpleDocTemplate(archivo, pagesize=pagesize)
            story = []
            styles = getSampleStyleSheet()

            # Título personalizado
            titulo_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1F4E78'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            titulo = Paragraph(f"📊 {self.informe_seleccionado}", titulo_style)
            story.append(titulo)

            # Información de fecha y filtros
            info_text = f"Generado: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
            if filtros_aplicados:
                info_text += f" | Filtros: {len(filtros_aplicados)}"
            if ordenaciones_aplicadas:
                info_text += f" | Ordenaciones: {len(ordenaciones_aplicadas)}"

            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey,
                alignment=TA_CENTER,
                spaceAfter=20
            )
            info = Paragraph(info_text, info_style)
            story.append(info)

            # Preparar datos para la tabla
            tabla_data = [columnas]  # Encabezados
            for fila in datos:
                tabla_data.append([str(v) if v is not None else "" for v in fila])

            # Crear tabla
            tabla = Table(tabla_data)

            # Estilo de la tabla
            tabla_style = TableStyle([
                # Encabezados
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),

                # Datos
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),

                # Bordes y líneas
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LINEABOVE', (0, 0), (-1, 0), 2, colors.HexColor('#4472C4')),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#4472C4')),

                # Alternar colores de fila
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')])
            ])

            tabla.setStyle(tabla_style)
            story.append(tabla)

            # Espacio
            story.append(Spacer(1, 0.3 * inch))

            # Resumen
            resumen_style = ParagraphStyle(
                'ResumenStyle',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_RIGHT,
                spaceAfter=12
            )
            resumen = Paragraph(f"<b>Total de registros: {len(datos)}</b>", resumen_style)
            story.append(resumen)

            # Generar PDF
            doc.build(story)

            CTkMessagebox(
                title="Exportación Exitosa",
                message=f"El informe se ha exportado correctamente a:\n\n{archivo}\n\n"
                        f"Registros exportados: {len(datos)}",
                icon="check"
            )

        except ImportError:
            CTkMessagebox(
                title="Error",
                message="La librería 'reportlab' no está instalada.\n\n"
                        "Por favor, instálala con:\n"
                        "pip install reportlab",
                icon="cancel"
            )
        except Exception as e:
            import traceback
            print(f"Error al exportar a PDF:\n{traceback.format_exc()}")
            CTkMessagebox(
                title="Error",
                message=f"Error al exportar a PDF:\n\n{str(e)}",
                icon="cancel"
            )

    def _export_message(self, format_name):
        """Muestra mensaje de exportación"""
        from CTkMessagebox import CTkMessagebox

        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Por favor, seleccione un informe del menú izquierdo.",
                icon="warning"
            )
            return

        CTkMessagebox(
            title=f"Exportar a {format_name}",
            message=f"Exportación a {format_name} en desarrollo.\n\n"
                    f"Informe seleccionado: {self.informe_seleccionado}",
            icon="info"
        )

    def _guardar_configuracion(self):
        """Guarda la configuración actual del informe"""
        from CTkMessagebox import CTkMessagebox
        import tkinter as tk
        
        # Validar que hay algo que guardar
        if not self.informe_seleccionado:
            CTkMessagebox(
                title="Aviso",
                message="Seleccione un informe primero.",
                icon="warning"
            )
            return
        
        # Recopilar configuración actual
        filtros_aplicados = self._recopilar_filtros()
        ordenaciones_aplicadas = self._recopilar_ordenaciones()
        campos_seleccionados_list = self._recopilar_campos()
        
        # Crear ventana de diálogo para nombrar la configuración
        dialog = customtkinter.CTkToplevel(self)
        dialog.title("Guardar Configuración")
        dialog.geometry("500x250")
        dialog.transient(self)
        dialog.grab_set()
        
        # Frame principal
        frame = customtkinter.CTkFrame(dialog)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        title_label = customtkinter.CTkLabel(
            frame,
            text="💾 Guardar Configuración de Informe",
            font=customtkinter.CTkFont(size=14, weight="bold")
        )
        title_label.pack(pady=(0, 15))
        
        # Nombre
        nombre_label = customtkinter.CTkLabel(frame, text="Nombre de la configuración:")
        nombre_label.pack(anchor="w", pady=(0, 5))
        
        nombre_entry = customtkinter.CTkEntry(frame, width=400, placeholder_text="Ej: Partes En Curso por OT")
        nombre_entry.pack(pady=(0, 10))
        nombre_entry.focus()
        
        # Descripción
        desc_label = customtkinter.CTkLabel(frame, text="Descripción (opcional):")
        desc_label.pack(anchor="w", pady=(0, 5))
        
        desc_entry = customtkinter.CTkEntry(frame, width=400, placeholder_text="Ej: Muestra partes en curso agrupados por OT")
        desc_entry.pack(pady=(0, 20))
        
        def guardar():
            nombre = nombre_entry.get().strip()
            if not nombre:
                CTkMessagebox(
                    title="Error",
                    message="El nombre es obligatorio.",
                    icon="cancel"
                )
                return
            
            descripcion = desc_entry.get().strip()
            
            # Guardar
            exito = self.storage.guardar_configuracion(
                nombre=nombre,
                informe_nombre=self.informe_seleccionado,
                filtros=filtros_aplicados,
                ordenaciones=ordenaciones_aplicadas,
                campos_seleccionados=campos_seleccionados_list,
                descripcion=descripcion
            )
            
            if exito:
                CTkMessagebox(
                    title="Éxito",
                    message=f"Configuración '{nombre}' guardada correctamente.",
                    icon="check"
                )
                dialog.destroy()
            else:
                CTkMessagebox(
                    title="Error",
                    message="No se pudo guardar la configuración.",
                    icon="cancel"
                )
        
        # Botones
        buttons_frame = customtkinter.CTkFrame(frame, fg_color="transparent")
        buttons_frame.pack()
        
        guardar_btn = customtkinter.CTkButton(
            buttons_frame,
            text="Guardar",
            width=120,
            command=guardar
        )
        guardar_btn.pack(side="left", padx=5)
        
        cancelar_btn = customtkinter.CTkButton(
            buttons_frame,
            text="Cancelar",
            width=120,
            fg_color="gray",
            command=dialog.destroy
        )
        cancelar_btn.pack(side="left", padx=5)
        
        # Centrar ventana
        dialog.update_idletasks()
        dialog.attributes('-topmost', True)
        dialog.lift()
        dialog.focus_force()
        dialog.after(100, lambda: dialog.attributes('-topmost', False))
    
    def _cargar_configuracion(self):
        """Carga una configuración guardada"""
        from CTkMessagebox import CTkMessagebox
        
        # Listar configuraciones disponibles
        configuraciones = self.storage.listar_configuraciones()
        
        if not configuraciones:
            CTkMessagebox(
                title="Aviso",
                message="No hay configuraciones guardadas.\n\nGuarde una configuración primero usando el botón '💾 Guardar Config'.",
                icon="info"
            )
            return
        
        # Crear ventana de diálogo para seleccionar configuración
        dialog = customtkinter.CTkToplevel(self)
        dialog.title("Cargar Configuración")
        dialog.geometry("700x500")
        dialog.transient(self)
        dialog.grab_set()
        
        # Frame principal
        frame = customtkinter.CTkFrame(dialog)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        title_label = customtkinter.CTkLabel(
            frame,
            text="📂 Cargar Configuración de Informe",
            font=customtkinter.CTkFont(size=14, weight="bold")
        )
        title_label.pack(pady=(0, 15))
        
        # Información
        info_label = customtkinter.CTkLabel(
            frame,
            text=f"Se encontraron {len(configuraciones)} configuraciones guardadas:",
            font=customtkinter.CTkFont(size=11)
        )
        info_label.pack(pady=(0, 10))
        
        # Frame scrollable para lista
        list_frame = customtkinter.CTkScrollableFrame(frame, height=300)
        list_frame.pack(fill="both", expand=True, pady=(0, 15))
        
        selected_config = {"name": None}
        
        def seleccionar(nombre):
            selected_config["name"] = nombre
            cargar()
        
        # Listar configuraciones
        for i, config in enumerate(configuraciones):
            config_frame = customtkinter.CTkFrame(list_frame)
            config_frame.pack(fill="x", pady=5, padx=5)
            
            # Info de la configuración
            info_text = f"📋 {config['nombre']}\n"
            info_text += f"   Informe: {config['informe_base']}\n"
            if config['descripcion']:
                info_text += f"   Descripción: {config['descripcion']}\n"
            info_text += f"   Filtros: {config['num_filtros']} | Ordenaciones: {config['num_ordenaciones']} | Campos: {config['num_campos']}\n"
            info_text += f"   Guardado: {config['fecha_creacion'][:10]}"
            
            label = customtkinter.CTkLabel(
                config_frame,
                text=info_text,
                justify="left",
                anchor="w"
            )
            label.pack(side="left", fill="x", expand=True, padx=10, pady=10)
            
            btn_frame = customtkinter.CTkFrame(config_frame, fg_color="transparent")
            btn_frame.pack(side="right", padx=10)
            
            cargar_btn = customtkinter.CTkButton(
                btn_frame,
                text="Cargar",
                width=80,
                command=lambda n=config['nombre']: seleccionar(n)
            )
            cargar_btn.pack(side="left", padx=2)
            
            eliminar_btn = customtkinter.CTkButton(
                btn_frame,
                text="🗑️",
                width=40,
                fg_color="darkred",
                hover_color="red",
                command=lambda n=config['nombre']: eliminar_config(n)
            )
            eliminar_btn.pack(side="left", padx=2)
        
        def cargar():
            nombre = selected_config["name"]
            if not nombre:
                return
            
            config = self.storage.cargar_configuracion(nombre)
            if not config:
                CTkMessagebox(
                    title="Error",
                    message="No se pudo cargar la configuración.",
                    icon="cancel"
                )
                return
            
            # Aplicar configuración
            self._aplicar_configuracion(config)
            
            CTkMessagebox(
                title="Éxito",
                message=f"Configuración '{nombre}' cargada correctamente.",
                icon="check"
            )
            dialog.destroy()
        
        def eliminar_config(nombre):
            respuesta = CTkMessagebox(
                title="Confirmar",
                message=f"¿Está seguro de eliminar la configuración '{nombre}'?",
                icon="question",
                option_1="Cancelar",
                option_2="Eliminar"
            )
            
            if respuesta.get() == "Eliminar":
                if self.storage.eliminar_configuracion(nombre):
                    CTkMessagebox(
                        title="Éxito",
                        message=f"Configuración '{nombre}' eliminada.",
                        icon="check"
                    )
                    dialog.destroy()
                    # Reabrir diálogo actualizado
                    self._cargar_configuracion()
        
        # Botón cerrar
        cancelar_btn = customtkinter.CTkButton(
            frame,
            text="Cerrar",
            width=120,
            fg_color="gray",
            command=dialog.destroy
        )
        cancelar_btn.pack()
        
        # Centrar ventana
        dialog.update_idletasks()
        dialog.attributes('-topmost', True)
        dialog.lift()
        dialog.focus_force()
        dialog.after(100, lambda: dialog.attributes('-topmost', False))
    
    def _recopilar_filtros(self):
        """Recopila los filtros actuales"""
        filtros_aplicados = []
        for filtro_obj in self.filtros:
            campo_actual = filtro_obj.get('campo_actual')
            if not campo_actual:
                continue
            
            operador = filtro_obj['operador_combo'].get()
            valor_widget = filtro_obj['valor_widget']
            
            # Obtener valor según tipo de widget
            if filtro_obj.get('is_range') and operador == "Entre":
                widget1 = filtro_obj.get('valor_widget')
                widget2 = filtro_obj.get('valor_widget2')
                
                if widget1 and widget2:
                    if isinstance(widget1, (customtkinter.CTkEntry, DateEntry)):
                        valor1 = widget1.get()
                    else:
                        valor1 = ""
                    
                    if isinstance(widget2, (customtkinter.CTkEntry, DateEntry)):
                        valor2 = widget2.get()
                    else:
                        valor2 = ""
                    
                    if not valor1 or not valor2:
                        continue
                    
                    valor = (valor1, valor2)
                else:
                    continue
            else:
                if isinstance(valor_widget, customtkinter.CTkComboBox):
                    valor = valor_widget.get()
                elif isinstance(valor_widget, (customtkinter.CTkEntry, DateEntry)):
                    valor = valor_widget.get()
                else:
                    valor = ""
            
            if not valor or valor == "Seleccionar..." or not operador or operador == "Seleccionar...":
                continue
            
            logica = 'Y'
            if filtro_obj.get('logica_combo'):
                logica = filtro_obj['logica_combo'].get()
            
            filtros_aplicados.append({
                'campo': campo_actual,
                'operador': operador,
                'valor': valor,
                'logica': logica
            })
        
        return filtros_aplicados
    
    def _recopilar_ordenaciones(self):
        """Recopila las ordenaciones actuales"""
        ordenaciones_aplicadas = []
        for clasif_obj in self.ordenaciones:
            campo_actual = clasif_obj.get('campo_actual')
            if not campo_actual:
                continue
            
            orden = clasif_obj['orden_combo'].get()
            if not orden or orden == "Seleccionar...":
                orden = "Ascendente"
            
            ordenaciones_aplicadas.append({
                'campo': campo_actual,
                'orden': orden
            })
        
        return ordenaciones_aplicadas
    
    def _recopilar_campos(self):
        """Recopila los campos seleccionados"""
        return [campo_key for campo_key, var in self.campos_seleccionados.items() if var.get()]
    
    def _aplicar_configuracion(self, config):
        """Aplica una configuración cargada"""
        # Limpiar estado actual
        self._clear_all_filtros()
        self._clear_all_ordenaciones()

        # Seleccionar informe base
        informe_base = config.get('informe_base')
        if informe_base:
            self.informe_seleccionado = informe_base
            self.definicion_actual = INFORMES_DEFINICIONES.get(informe_base)

            # Actualizar título
            if self.definicion_actual:
                titulo = f"Informe seleccionado: {informe_base}"
                descripcion = self.definicion_actual.get('descripcion', '')
                if descripcion:
                    titulo += f" ({descripcion})"
                self.informe_title_label.configure(text=titulo, text_color="white")

        # Cargar campos disponibles
        self._update_campos_disponibles()

        # Aplicar campos seleccionados
        campos_config = config.get('campos_seleccionados', [])
        for campo_key, var in self.campos_seleccionados.items():
            var.set(campo_key in campos_config)

        # Aplicar filtros
        filtros_config = config.get('filtros', [])
        for i, filtro_data in enumerate(filtros_config):
            self._add_filtro()
            if i < len(self.filtros):
                self._configurar_filtro(self.filtros[i], filtro_data)

        # Aplicar ordenaciones
        clasifs_config = config.get('ordenaciones', [])
        for i, clasif_data in enumerate(clasifs_config):
            self._add_ordenacion()
            if i < len(self.ordenaciones):
                self._configurar_ordenacion(self.ordenaciones[i], clasif_data)

    def _configurar_filtro(self, filtro_obj, filtro_data):
        """Configura un filtro con los datos guardados"""
        if not self.definicion_actual:
            return

        campos_def = self.definicion_actual.get('campos', {})

        # 1. Configurar lógica (Y/O) si no es el primero
        if filtro_obj['logica_combo'] and 'logica' in filtro_data:
            filtro_obj['logica_combo'].set(filtro_data['logica'])

        # 2. Configurar campo
        campo_key = filtro_data.get('campo')
        if campo_key and campo_key in campos_def:
            campo_nombre = campos_def[campo_key]['nombre']
            filtro_obj['campo_combo'].set(campo_nombre)
            self._on_filtro_campo_change(filtro_obj, campo_nombre)

            # 3. Configurar operador
            operador = filtro_data.get('operador')
            if operador:
                filtro_obj['operador_combo'].set(operador)
                self._on_filtro_operador_change(filtro_obj, operador)

                # 4. Configurar valor
                valor = filtro_data.get('valor')
                if valor is not None:
                    # Esperar un momento para que se cree el widget
                    self.after(100, lambda: self._set_filtro_valor(filtro_obj, operador, valor))

    def _set_filtro_valor(self, filtro_obj, operador, valor):
        """Establece el valor de un filtro según su tipo"""
        if operador == "Entre" and isinstance(valor, (list, tuple)) and len(valor) == 2:
            # Valor de rango (min, max)
            widget1 = filtro_obj.get('valor_widget')
            widget2 = filtro_obj.get('valor_widget2')

            if widget1 and widget2:
                if isinstance(widget1, DateEntry):
                    # Para DateEntry, usar set_date
                    try:
                        from datetime import datetime
                        fecha1 = datetime.strptime(valor[0], '%Y-%m-%d')
                        widget1.set_date(fecha1)
                    except:
                        pass
                else:
                    # Para Entry, usar insert
                    widget1.delete(0, 'end')
                    widget1.insert(0, str(valor[0]))

                if isinstance(widget2, DateEntry):
                    try:
                        from datetime import datetime
                        fecha2 = datetime.strptime(valor[1], '%Y-%m-%d')
                        widget2.set_date(fecha2)
                    except:
                        pass
                else:
                    widget2.delete(0, 'end')
                    widget2.insert(0, str(valor[1]))
        else:
            # Valor simple
            widget = filtro_obj.get('valor_widget')
            if widget:
                if isinstance(widget, customtkinter.CTkComboBox):
                    widget.set(str(valor))
                elif isinstance(widget, DateEntry):
                    try:
                        from datetime import datetime
                        fecha = datetime.strptime(valor, '%Y-%m-%d')
                        widget.set_date(fecha)
                    except:
                        pass
                elif isinstance(widget, customtkinter.CTkEntry):
                    widget.delete(0, 'end')
                    widget.insert(0, str(valor))

    def _configurar_ordenacion(self, clasif_obj, clasif_data):
        """Configura una ordenación con los datos guardados"""
        if not self.definicion_actual:
            return

        campos_def = self.definicion_actual.get('campos', {})

        # 1. Configurar campo
        campo_key = clasif_data.get('campo')
        if campo_key and campo_key in campos_def:
            campo_nombre = campos_def[campo_key]['nombre']
            clasif_obj['var_combo'].set(campo_nombre)
            self._on_ordenacion_campo_change(clasif_obj, campo_nombre)

        # 2. Configurar orden
        orden = clasif_data.get('orden', 'Ascendente')
        if clasif_obj['orden_combo']:
            clasif_obj['orden_combo'].set(orden)
